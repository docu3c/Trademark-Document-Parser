# Version - 6.1  (Code Updated of Checking for Non-matching Class number == GPT 4o mini)

from fileinput import filename
import time, os
import streamlit as st
import pandas as pd
import fitz  # PyMuPDF
from pydantic import BaseModel, Field, ValidationError
from typing import List, Dict, Union
import base64
from docx import Document
from docx.shared import Pt
from io import BytesIO
import re, ast
from dotenv import load_dotenv

load_dotenv()

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


class TrademarkDetails(BaseModel):
    trademark_name: str = Field(
        description="The name of the Trademark", example="DISCOVER"
    )
    status: str = Field(description="The Status of the Trademark", example="Registered")
    serial_number: str = Field(
        description="The Serial Number of the trademark from Chronology section",
        example="87−693,628",
    )
    international_class_number: List[int] = Field(
        description="The International class number or Nice Classes number of the trademark from Goods/Services section or Nice Classes section",
        example=[18],
    )
    owner: str = Field(
        description="The owner of the trademark", example="WALMART STORES INC"
    )
    goods_services: str = Field(
        description="The goods/services from the document",
        example="LUGGAGE AND CARRYING BAGS; SUITCASES, TRUNKS, TRAVELLING BAGS, SLING BAGS FOR CARRYING INFANTS, SCHOOL BAGS; PURSES; WALLETS; RETAIL AND ONLINE RETAIL SERVICES",
    )
    page_number: int = Field(
        description="The page number where the trademark details are found in the document",
        example=3,
    )
    registration_number: Union[str, None] = Field(
        description="The Registration number of the trademark from Chronology section",
        example="5,809,957",
    )
    design_phrase: str = Field(
        description="The design phrase of the trademark",
        example="THE MARK CONSISTS OF THE STYLIZED WORD 'MINI' FOLLOWED BY 'BY MOTHERHOOD.'",
        default="",
    )


# azure_endpoint = st.secrets["AZURE_ENDPOINT"]
# api_key = st.secrets["AZURE_API_KEY"]


def preprocess_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"[\u2013\u2014]", "-", text)
    return text


def is_correct_format_code1(page_text: str) -> bool:
    required_fields = ["Status:", "Goods/Services:"]  # , "Last Reported Owner:"
    return all(field in page_text for field in required_fields)


def is_correct_format_code2(page_text: str) -> bool:
    required_fields = ["Register", "Nice Classes", "Goods & Services"]
    return all(field in page_text for field in required_fields)


def extract_trademark_details_code1(
    document_chunk: str,
) -> Dict[str, Union[str, List[int]]]:
    try:
        from openai import AzureOpenAI

        azure_endpoint = os.getenv("AZURE_ENDPOINT")
        api_key = os.getenv("AZURE_API_KEY")

        client = AzureOpenAI(
            azure_endpoint=azure_endpoint,
            api_key=api_key,
            api_version="2024-08-01-preview",
        )

        messages = [
            {
                "role": "system",
                "content": "You are a helpful assistant for extracting Meta Data from the Trademark Document.",
            },
            {
                "role": "user",
                "content": f"""
                Extract the following details from the trademark document: trademark name, status.\n\nDocument:\n{document_chunk}
                Don't extract the same trademark details more than once; extract them only once. 
                 
                Return output only in the below mentioned format:
                Example-1 output format: 
                    Trademark Name: SLIK\n 
                    Status: PENDING\n
                Example-2 output format: 
                    Trademark Name: HUMOR US GOODS\n 
                    Status: REGISTERED\n
                Example-3 output format: 
                    Trademark Name: #WASONUO %& PIC\n 
                    Status: REGISTERED\n
                Example-4 output format: 
                    Trademark Name: AT Present, WE’VE GOT YOUR-BACK(SIDE)\n 
                    Status: PUBLISHED\n\n
                    
                Note: The trademark name length can also be 1 or 2 characters. (Example: Trademark Name: PI), (Example: Trademark Name: PII) \n"""
            },
        ]
        
                # Example-5 output format: 
                #     Trademark Name: PI\n
                #     Status: REGISTERED\n
                    
        # Not available in the provided document
        #  Example expected output format: Trademark Name: SLIK Status: PENDING FILED AS USE APPLICATION Serial Number: 98-602,112 International Class Number: 3 Owner: SLIK DE VENEZUELA C.A. VENEZUELA CORPORATION Goods & Services: Cosmetics; hair gel; hair wax; hair styling gel; non-medicated cosmetics Filed Date: JUN 14, 2024 Registration Number: Not available in the provided document.
        #  Example expected output: Trademark Name: #WASONOFILTER Status: REGISTERED Serial Number: 88-404,432 International Class Number: 21 Owner: LAHC US 1 LLC DELAWARE LIMITED LIABILITY COMPANY Goods & Services:  Containers for household use, coffee mugs, and wine glasses Filed Date: APR 26, 2019 Registration Number: 5,963,355"""}

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            temperature=0,
            max_tokens=300,
        )
        extracted_text = response.choices[0].message.content

        # if extracted_text and extracted_text != "[]":
        # st.write(extracted_text)

        details = {}
        for line in extracted_text.split("\n"):
            if ":" in line:
                key, value = line.split(":", 1)
                details[key.strip().lower().replace(" ", "_")] = value.strip()

        # st.warning(details)
        return details

    except Exception as e:
        print(f"An error occurred: {e}")


def extract_serial_number(
    document: str, start_page: int, pdf_document: fitz.Document
) -> str:
    combined_texts = ""
    for i in range(start_page, min(start_page + 13, pdf_document.page_count)):
        page = pdf_document.load_page(i)
        page_text = page.get_text()
        combined_texts += page_text
        if "Serial Number:" in page_text or "Ownership Details:" in page_text:
            break

    pattern = r"Chronology:.*?Serial Number:\s*([\d,-−]+)"
    match = re.search(pattern, combined_texts, re.DOTALL)
    if match:
        registration_number = match.group(1).strip()
        return registration_number
    return "No serial number presented in document"


def extract_ownership(document: str, start_page: int, proposed_name: str, pdf_document: fitz.Document) -> str:
    """ Extract the ownership from the document """
    combined_texts = ""
    for i in range(start_page, min(start_page + 13, pdf_document.page_count)):
        page = pdf_document.load_page(i)
        page_text = page.get_text()
        combined_texts += page_text
        if "Last Reported Owner:" in page_text or "Ownership Details:" in page_text:
            break

    pattern = r"Last Reported Owner:\s*(.*?)\n\s*(.*?)\n"
    match = re.search(pattern, combined_texts, re.DOTALL)
    if match:
        owner_name = match.group(1).strip()
        owner_type = match.group(2).strip()
        if owner_type == proposed_name:
            return f"{owner_name}"
        else:
            return f"{owner_name} {owner_type}"
    return "Not available in the provided document."


def extract_registration_number(
    document: str, start_page: int, pdf_document: fitz.Document
) -> str:
    """Extract the registration number from the document"""
    combined_texts = ""
    for i in range(start_page, min(start_page + 8, pdf_document.page_count)):
        page = pdf_document.load_page(i)
        page_text = page.get_text()
        combined_texts += page_text
        if "Registration Number:" in page_text or "Ownership Details:" in page_text:
            break

    pattern = r"Last ReportedOwner:.*?Registration Number:\s*([\d,]+)"
    match = re.search(pattern, combined_texts, re.DOTALL)
    if match:
        registration_number = match.group(1).strip()
        return registration_number
    return "NA"


def extract_trademark_details_code2(page_text: str) -> Dict[str, Union[str, List[int]]]:
    details = {}

    trademark_name_match = re.search(
        r"\d+\s*/\s*\d+\s*\n\s*\n\s*([A-Za-z0-9'&!,\-. ]+)\s*\n", page_text
    )
    if trademark_name_match:
        details["trademark_name"] = trademark_name_match.group(1).strip()
    else:
        trademark_name_match = re.search(
            r"(?<=\n)([A-Za-z0-9'&!,\-. ]+)(?=\n)", page_text
        )
        details["trademark_name"] = (
            trademark_name_match.group(1).strip() if trademark_name_match else ""
        )

    status_match = re.search(
        r"Status\s*(?:\n|:\s*)([A-Za-z]+)", page_text, re.IGNORECASE
    )
    details["status"] = status_match.group(1).strip() if status_match else ""

    owner_match = re.search(r"Holder\s*(?:\n|:\s*)(.*)", page_text, re.IGNORECASE)
    if owner_match:
        details["owner"] = owner_match.group(1).strip()
    else:
        owner_match = re.search(r"Owner\s*(?:\n|:\s*)(.*)", page_text, re.IGNORECASE)
        details["owner"] = owner_match.group(1).strip() if owner_match else ""

    nice_classes_match = re.search(
        r"Nice Classes\s*[\s:]*\n((?:\d+(?:,\s*\d+)*)\b)", page_text, re.IGNORECASE
    )
    if nice_classes_match:
        nice_classes_text = nice_classes_match.group(1)
        nice_classes = [int(cls.strip()) for cls in nice_classes_text.split(",")]
        details["international_class_number"] = nice_classes
    else:
        details["international_class_number"] = []

    serial_number_match = re.search(r"Application#\s*(.*)", page_text, re.IGNORECASE)
    details["serial_number"] = (
        serial_number_match.group(1).strip() if serial_number_match else ""
    )

    goods_services_match = re.search(
        r"Goods & Services\s*(.*?)(?=\s*G&S translation|$)",
        page_text,
        re.IGNORECASE | re.DOTALL,
    )
    details["goods_services"] = (
        goods_services_match.group(1).strip() if goods_services_match else ""
    )

    registration_number_match = re.search(
        r"Registration#\s*(.*)", page_text, re.IGNORECASE
    )
    details["registration_number"] = (
        registration_number_match.group(1).strip() if registration_number_match else ""
    )

    # Description
    design_phrase = re.search(
        r"Description\s*(.*?)(?=\s*Applicant|Owner|Holder|$)",
        page_text,
        re.IGNORECASE | re.DOTALL,
    )
    details["design_phrase"] = (
        design_phrase.group(1).strip()
        if design_phrase
        else "No Design phrase presented in document"
    )

    return details


def read_pdf(file_path: str, exclude_header_footer: bool = True) -> str:
    document_text = ""
    with fitz.open(file_path) as pdf_document:
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            if exclude_header_footer:
                rect = page.rect
                x0 = rect.x0
                y0 = rect.y0 + rect.height * 0.1
                x1 = rect.x1
                y1 = rect.y1 - rect.height * 0.1
                page_text = page.get_text("text", clip=(x0, y0, x1, y1))
            else:
                page_text = page.get_text()
            document_text += page_text
    return document_text


def split_text(text: str, max_tokens: int = 1500) -> List[str]:
    chunks = []
    current_chunk = []
    current_length = 0

    for line in text.split("\n"):
        line_length = len(line.split())
        if current_length + line_length > max_tokens:
            chunks.append("\n".join(current_chunk))
            current_chunk = [line]
            current_length = line_length
        else:
            current_chunk.append(line)
            current_length += line_length

    if current_chunk:
        chunks.append("\n".join(current_chunk))

    return chunks


def parse_international_class_numbers(class_numbers: str) -> List[int]:
    numbers = class_numbers.split(",")
    return [int(num.strip()) for num in numbers if num.strip().isdigit()]


def extract_international_class_numbers_and_goods_services(
    document: str, start_page: int, pdf_document: fitz.Document
) -> Dict[str, Union[List[int], str]]:
    """Extract the International Class Numbers and Goods/Services from the document over a range of pages"""
    class_numbers = []
    goods_services = []
    combined_text = ""

    for i in range(start_page, min(start_page + 10, pdf_document.page_count)):
        page = pdf_document.load_page(i)
        page_text = page.get_text()
        combined_text += page_text
        if "Last Reported Owner:" in page_text:
            break

    pattern = r"International Class (\d+): (.*?)(?=\nInternational Class \d+:|\n[A-Z][a-z]+:|\nLast Reported Owner:|Disclaimers:|\Z)"
    matches = re.findall(pattern, combined_text, re.DOTALL)
    for match in matches:
        class_number = int(match[0])
        class_numbers.append(class_number)
        goods_services.append(f"Class {class_number}: {match[1].strip()}")

    if "sexual" in goods_services or "sex" in goods_services:
        goods_services = replace_disallowed_words(goods_services)

    return {
        "international_class_numbers": class_numbers,
        "goods_services": "\n".join(goods_services),
    }


def extract_design_phrase(
    document: str, start_page: int, pdf_document: fitz.Document
) -> str:
    """Extract the design phrase from the document"""
    combined_texts = ""
    for i in range(start_page, min(start_page + 10, pdf_document.page_count)):
        page = pdf_document.load_page(i)
        page_text = page.get_text()
        combined_texts += page_text
        if "Design Phrase:" in page_text or "Filing Correspondent:" in page_text:
            break

    pattern = r"Design Phrase:\s*(.*?)(?=Other U\.S\. Registrations:|Filing Correspondent:|Group:|USPTO Page:|$)"
    match = re.search(pattern, combined_texts, re.DOTALL)
    if match:
        design_phrase = match.group(1).strip()
        # Remove any newline characters within the design phrase
        design_phrase = " ".join(design_phrase.split())
        return design_phrase
    return "No Design phrase presented in document"


def parse_trademark_details(
    document_path: str,
) -> List[Dict[str, Union[str, List[int]]]]:
    with fitz.open(document_path) as pdf_document:
        all_extracted_data = []
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            page_text = page.get_text()

            if is_correct_format_code1(page_text):
                preprocessed_chunk = preprocess_text(page_text)
                extracted_data = extract_trademark_details_code1(preprocessed_chunk)
                additional_data = (
                    extract_international_class_numbers_and_goods_services(
                        page_text, page_num, pdf_document
                    )
                )
                registration_number = extract_registration_number(
                    page_text, page_num, pdf_document
                )
                serial_number = extract_serial_number(page_text, page_num, pdf_document)
                design_phrase = extract_design_phrase(page_text, page_num, pdf_document)
                ownership_details = extract_ownership(page_text, page_num, proposed_name, pdf_document)
                

                if extracted_data:
                    extracted_data["page_number"] = page_num + 1
                    extracted_data.update(additional_data)
                    extracted_data["design_phrase"] = design_phrase
                    extracted_data["owner"] = ownership_details
                    extracted_data["serial_number"] = serial_number
                    extracted_data["registration_number"] = registration_number
                    all_extracted_data.append(extracted_data)

                trademark_list = []
                for i, data in enumerate(all_extracted_data, start=1):
                    try:                        
                        trademark_name = data.get("trademark_name", "")
                        if "Global Filings" in trademark_name:
                            trademark_name = trademark_name.split("Global Filings")[
                                0
                            ].strip()
                        if re.match(r"^US-\d+", trademark_name):
                            trademark_name = re.sub(r"^US-\d+\s*", "", trademark_name).strip()
                        status = data.get("status", "").split(",")[0].strip()
                        serial_number = data.get("serial_number", "")
                        owner = data.get("owner", "")
                        international_class_number = data.get(
                            "international_class_numbers", []
                        )
                        goods_services = data.get("goods_services", "")
                        page_number = data.get("page_number", "")
                        registration_number = data.get(
                            "registration_number",
                            "No registration number presented in document",
                        )
                        design_phrase = data.get(
                            "design_phrase", "No Design phrase presented in document"
                        )

                        # If crucial fields are missing, attempt to re-extract the values

                        # if not trademark_name or not owner or not status or not international_class_number:
                        #     preprocessed_chunk = preprocess_text(data.get("raw_text", ""))
                        #     extracted_data = extract_trademark_details_code1(preprocessed_chunk)
                        #     trademark_name = extracted_data.get("trademark_name", trademark_name).split(',')[0].strip()
                        #     if "Global Filings" in trademark_name:
                        #         trademark_name = trademark_name.split("Global Filings")[0].strip()
                        #     owner = extracted_data.get("owner", owner).split(',')[0].strip()
                        #     status = extracted_data.get("status", status).split(',')[0].strip()
                        #     international_class_number = parse_international_class_numbers(extracted_data.get("international_class_number", "")) or international_class_number
                        #     registration_number = extracted_data.get("registration_number", registration_number).split(',')[0].strip()

                        trademark_details = TrademarkDetails(
                            trademark_name=trademark_name,
                            owner=owner,
                            status=status,
                            serial_number=serial_number,
                            international_class_number=international_class_number,
                            goods_services=goods_services,
                            page_number=page_number,
                            registration_number=registration_number,
                            design_phrase=design_phrase,
                        )
                        trademark_info = {
                            "trademark_name": trademark_details.trademark_name,
                            "owner": trademark_details.owner,
                            "status": trademark_details.status,
                            "serial_number": trademark_details.serial_number,
                            "international_class_number": trademark_details.international_class_number,
                            "goods_services": trademark_details.goods_services,
                            "page_number": trademark_details.page_number,
                            "registration_number": trademark_details.registration_number,
                            "design_phrase": trademark_details.design_phrase,
                        }
                        print(trademark_info)
                        print(
                            "_____________________________________________________________________________________________________________________________"
                        )
                        trademark_list.append(trademark_info)
                    except ValidationError as e:
                        print(f"Validation error for trademark {i}: {e}")

            else:
                if not is_correct_format_code2(page_text):
                    continue

                extracted_data = extract_trademark_details_code2(page_text)
                st.info("Code 2")
                if extracted_data:
                    extracted_data["page_number"] = page_num + 1
                    all_extracted_data.append(extracted_data)

                trademark_list = []
                for i, data in enumerate(all_extracted_data, start=1):
                    try:
                        trademark_details = TrademarkDetails(
                            trademark_name=data.get("trademark_name", ""),
                            owner=data.get("owner", ""),
                            status=data.get("status", ""),
                            serial_number=data.get("serial_number", ""),
                            international_class_number=data.get(
                                "international_class_number", []
                            ),
                            goods_services=data.get("goods_services", ""),
                            page_number=data.get("page_number", 0),
                            registration_number=data.get("registration_number", ""),
                            design_phrase=data.get("design_phrase", ""),
                        )
                        if (
                            trademark_details.trademark_name != ""
                            and trademark_details.owner != ""
                            and trademark_details.status != ""
                            and trademark_details.goods_services != ""
                        ):
                            trademark_info = {
                                "trademark_name": trademark_details.trademark_name,
                                "owner": trademark_details.owner,
                                "status": trademark_details.status,
                                "serial_number": trademark_details.serial_number,
                                "international_class_number": trademark_details.international_class_number,
                                "goods_services": trademark_details.goods_services,
                                "page_number": trademark_details.page_number,
                                "registration_number": trademark_details.registration_number,
                                "design_phrase": trademark_details.design_phrase,
                            }

                            trademark_list.append(trademark_info)
                    except ValidationError as e:
                        print(f"Validation error for trademark {i}: {e}")

        return trademark_list


from typing import List, Dict, Union
from sentence_transformers import SentenceTransformer, util
from fuzzywuzzy import fuzz

# Load the semantic similarity model
semantic_model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")


def compare_trademarks(
    existing_trademark: Dict[str, Union[str, List[int]]],
    proposed_name: str,
    proposed_class: str,
    proposed_goods_services: str,
) -> Dict[str, Union[str, int]]:
    # Convert proposed classes to a list of integers
    proposed_classes = [int(c.strip()) for c in proposed_class.split(",")]

    # Helper function for semantic equivalence
    def is_semantically_equivalent(name1, name2, threshold=0.80):
        embeddings1 = semantic_model.encode(name1, convert_to_tensor=True)
        embeddings2 = semantic_model.encode(name2, convert_to_tensor=True)
        similarity_score = util.cos_sim(embeddings1, embeddings2).item()
        return similarity_score >= threshold

    # Helper function for phonetic equivalence
    def is_phonetically_equivalent(name1, name2, threshold=80):
        return fuzz.ratio(name1.lower(), name2.lower()) >= threshold

    # Helper function for phonetically equivalent words
    def first_words_phonetically_equivalent(existing_name, proposed_name, threshold=80):
        existing_words = existing_name.lower().split()
        proposed_words = proposed_name.lower().split()
        if len(existing_words) < 2 or len(proposed_words) < 2:
            return False
        return (
            fuzz.ratio(" ".join(existing_words[:2]), " ".join(proposed_words[:2]))
            >= threshold
        )

    # def is_exact_match(name1: str, name2: str) -> bool:
    #     # Initial exact match check
    #     if name1.strip().lower() == name2.strip().lower():
    #         return True
    #     else:
    #         # Check for near-exact matches using normalized forms
    #         normalized_name1 = normalize_texts(name1)
    #         normalized_name2 = normalize_texts(name2)
    #         if normalized_name1 == normalized_name2:
    #             return True
    #         elif fuzz.ratio(normalized_name1, normalized_name2) >= 95:
    #             # Near-exact match, supplement with LLM
    #             return is_exact_match_llm(name1, name2)
    #         else:
    #             return False

    # def normalize_texts(text: str) -> str:
    #     import unicodedata
    #     import re

    #     # Normalize unicode characters
    #     text = unicodedata.normalize("NFKD", text)
    #     # Remove diacritics
    #     text = "".join(c for c in text if not unicodedata.combining(c))
    #     # Remove special characters and punctuation
    #     text = re.sub(r"[^\w\s]", "", text)
    #     # Convert to lowercase and strip whitespace
    #     return text.lower().strip()

    # def is_exact_match_llm(name1: str, name2: str) -> bool:
    #     from openai import AzureOpenAI
    #     import os

    #     azure_endpoint = os.getenv("AZURE_ENDPOINT")
    #     api_key = os.getenv("AZURE_API_KEY")
    #     client = AzureOpenAI(
    #         azure_endpoint=azure_endpoint,
    #         api_key=api_key,
    #         api_version="2024-10-01-preview",
    #     )

    #     prompt = f"""  
    #         Are the following two trademark names considered exact matches, accounting for minor variations such as special characters, punctuation, or formatting? Respond with 'Yes' or 'No'.  
            
    #         Trademark Name 1: "{name1}"  
    #         Trademark Name 2: "{name2}"  
    #         """

    #     messages = [
    #         {
    #             "role": "system",
    #             "content": "You are a trademark expert specializing in name comparisons.",
    #         },
    #         {"role": "user", "content": prompt},
    #     ]

    #     response = client.chat.completions.create(
    #         model="gpt-4o-mini",
    #         messages=messages,
    #         temperature=0.0,
    #         max_tokens=5,
    #     )

    #     answer = response.choices[0].message.content.strip().lower()
    #     return "yes" in answer.lower()

    # def is_semantically_equivalents(
    #     name1: str, name2: str, threshold: float = 0.80
    # ) -> bool:
    #     embeddings1 = semantic_model.encode(name1, convert_to_tensor=True)
    #     embeddings2 = semantic_model.encode(name2, convert_to_tensor=True)
    #     similarity_score = util.cos_sim(embeddings1, embeddings2).item()
    #     if similarity_score >= threshold:
    #         return True
    #     elif similarity_score >= (threshold - 0.1):
    #         # Near-threshold case, supplement with LLM
    #         return is_semantically_equivalent_llm(name1, name2)
    #     else:
    #         return False

    # def is_semantically_equivalent_llm(name1: str, name2: str) -> bool:
    #     prompt = f"""  
    #     Are the following two trademark names semantically equivalent? Respond with 'Yes' or 'No'.  
        
    #     Trademark Name 1: "{name1}"  
    #     Trademark Name 2: "{name2}"  
    #     """

    #     azure_endpoint = os.getenv("AZURE_ENDPOINT")
    #     api_key = os.getenv("AZURE_API_KEY")
    #     client = AzureOpenAI(
    #         azure_endpoint=azure_endpoint,
    #         api_key=api_key,
    #         api_version="2024-10-01-preview",
    #     )

    #     messages = [
    #         {
    #             "role": "system",
    #             "content": "You are an expert in trademark law and semantics.",
    #         },
    #         {"role": "user", "content": prompt},
    #     ]

    #     response = client.chat.completions.create(
    #         model="gpt-4o-mini",
    #         messages=messages,
    #         temperature=0.0,
    #         max_tokens=5,
    #     )

    #     answer = response.choices[0].message.content.strip().lower()
    #     return "yes" in answer.lower()

    # def is_phonetically_equivalents(
    #     name1: str, name2: str, threshold: int = 80
    # ) -> bool:
    #     from metaphone import doublemetaphone

    #     dm_name1 = doublemetaphone(name1)
    #     dm_name2 = doublemetaphone(name2)
    #     phonetic_similarity = fuzz.ratio(dm_name1[0], dm_name2[0])
    #     if phonetic_similarity >= threshold:
    #         return True
    #     elif phonetic_similarity >= (threshold - 10):
    #         # Near-threshold case, supplement with LLM
    #         return is_phonetically_equivalent_llm(name1, name2)
    #     else:
    #         return False

    # def is_phonetically_equivalent_llm(name1: str, name2: str) -> bool:

    #     prompt = f"""  
    #     Do the following two trademark names sound the same or very similar when spoken aloud? Consider differences in spelling but similarities in pronunciation. Respond with 'Yes' or 'No'.  
        
    #     Trademark Name 1: "{name1}"  
    #     Trademark Name 2: "{name2}"  
    #     """

    #     messages = [
    #         {
    #             "role": "system",
    #             "content": "You are an expert in phonetics and trademark law.",
    #         },
    #         {"role": "user", "content": prompt},
    #     ]

    #     azure_endpoint = os.getenv("AZURE_ENDPOINT")
    #     api_key = os.getenv("AZURE_API_KEY")
    #     client = AzureOpenAI(
    #         azure_endpoint=azure_endpoint,
    #         api_key=api_key,
    #         api_version="2024-10-01-preview",
    #     )

    #     response = client.chat.completions.create(
    #         model="gpt-4o-mini",
    #         messages=messages,
    #         temperature=0.0,
    #         max_tokens=5,
    #     )

    #     answer = response.choices[0].message.content.strip().lower()
    #     return "yes" in answer.lower()

    # Condition 1A: Exact character-for-character match
    condition_1A_satisfied = (
        existing_trademark["trademark_name"].strip().lower()
        == proposed_name.strip().lower()
    )

    # Condition 1B: Semantically equivalent
    condition_1B_satisfied = is_semantically_equivalent(
        existing_trademark["trademark_name"], proposed_name
    )

    # Condition 1C: Phonetically equivalent
    condition_1C_satisfied = is_phonetically_equivalent(
        existing_trademark["trademark_name"], proposed_name
    )

    # Condition 1D: First two or more words are phonetically equivalent
    condition_1D_satisfied = first_words_phonetically_equivalent(
        existing_trademark["trademark_name"], proposed_name
    )

    # Condition 1E: Proposed name is the first word of the existing trademark
    condition_1E_satisfied = (
        existing_trademark["trademark_name"].lower().startswith(proposed_name.lower())
    )

    # Check if any Condition 1 is satisfied
    condition_1_satisfied = any(
        [
            condition_1A_satisfied,
            condition_1B_satisfied,
            condition_1C_satisfied,
            condition_1D_satisfied,
            condition_1E_satisfied,
        ]
    )

    # def target_market_and_goods_overlaps(existing_gs, proposed_gs, threshold=0.65):
    #     embeddings1 = semantic_model.encode(existing_gs, convert_to_tensor=True)
    #     embeddings2 = semantic_model.encode(proposed_gs, convert_to_tensor=True)
    #     similarity_score = util.cos_sim(embeddings1, embeddings2).item()
    #     if similarity_score >= threshold:
    #         return True
    #     elif similarity_score >= (threshold - 0.1):
    #         # Supplement with LLM
    #         return target_market_and_goods_overlap_llm(existing_gs, proposed_gs)
    #     else:
    #         # Further check using keyword overlap
    #         # ... Additional code
    #         return False

    # def target_market_and_goods_overlap_llm(existing_gs: str, proposed_gs: str) -> bool:
    #     prompt = f"""  
    #         Do the goods and services described in the existing trademark and the proposed trademark overlap or target the same market? Consider the descriptions carefully. Respond with 'Yes' or 'No'.  
            
    #         Existing Trademark Goods/Services:  
    #         "{existing_gs}"  
            
    #         Proposed Trademark Goods/Services:  
    #         "{proposed_gs}"  
    #         """

    #     messages = [
    #         {
    #             "role": "system",
    #             "content": "You are an expert in trademark law and market analysis.",
    #         },
    #         {"role": "user", "content": prompt},
    #     ]

    #     azure_endpoint = os.getenv("AZURE_ENDPOINT")
    #     api_key = os.getenv("AZURE_API_KEY")
    #     client = AzureOpenAI(
    #         azure_endpoint=azure_endpoint,
    #         api_key=api_key,
    #         api_version="2024-10-01-preview",
    #     )

    #     response = client.chat.completions.create(
    #         model="gpt-4o-mini",
    #         messages=messages,
    #         temperature=0.0,
    #         max_tokens=5,
    #     )

    #     answer = response.choices[0].message.content.strip().lower()
    #     return "yes" in answer.lower()

    # Condition 2: Overlap in International Class Numbers
    condition_2_satisfied = bool(
        set(existing_trademark["international_class_number"]) & set(proposed_classes)
    )

    import re
    from nltk.stem import WordNetLemmatizer

    def normalize_text(text):

        # Replace special hyphen-like characters with a standard hyphen
        text = re.sub(r"[−–—]", "-", text)
        # Remove punctuation except hyphens and spaces
        text = re.sub(r"[^\w\s-]", " ", text)
        # Convert to lowercase
        text = text.lower()
        text = re.sub(r"\b\d+\b", "", text)
        text = re.sub(r"\bclass\b", "", text)
        text = re.sub(r"\bcare\b", "", text)
        text = re.sub(r"\bin\b", "", text)
        text = re.sub(r"\band\b", "", text)
        text = re.sub(r"\bthe\b", "", text)
        text = re.sub(r"\bfor\b", "", text)
        text = re.sub(r"\bwith\b", "", text)
        text = re.sub(r"\bfrom\b", "", text)
        text = re.sub(r"\bto\b", "", text)
        text = re.sub(r"\bunder\b", "", text)
        text = re.sub(r"\busing\b", "", text)
        text = re.sub(r"\bof\b", "", text)
        text = re.sub(r"\bno\b", "", text)
        text = re.sub(r"\binclude\b", "", text)
        text = re.sub(r"\bex\b", "", text)
        text = re.sub(r"\bexample\b", "", text)
        text = re.sub(r"\bclasses\b", "", text)
        text = re.sub(r"\bsearch\b", "", text)
        text = re.sub(r"\bscope\b", "", text)
        text = re.sub(r"\bshower\b", "", text)
        text = re.sub(r"\bproducts\b", "", text)
        text = re.sub(r"\bshampoos\b", "hair", text)

        # Standardize whitespace
        return " ".join(text.split())

    # Condition 3: Target market and goods/services overlap
    def target_market_and_goods_overlap(existing_gs, proposed_gs, threshold=0.65):

        existing_normalized = normalize_text(existing_gs)
        proposed_normalized = normalize_text(proposed_gs)

        embeddings1 = semantic_model.encode(existing_normalized, convert_to_tensor=True)
        embeddings2 = semantic_model.encode(proposed_normalized, convert_to_tensor=True)
        similarity_score = util.cos_sim(embeddings1, embeddings2).item()
        # st.write("Semantic Similarity Score:", similarity_score)
        if similarity_score >= threshold:
            return True

        # Split into words and lemmatize
        lemmatizer = WordNetLemmatizer()
        existing_words = {
            lemmatizer.lemmatize(word) for word in existing_normalized.split()
        }
        proposed_words = {
            lemmatizer.lemmatize(word) for word in proposed_normalized.split()
        }

        # Check for common words
        common_words = existing_words.intersection(proposed_words)
        # st.write("Common Words:", existing_gs , common_words)
        return bool(common_words)

    condition_3_satisfied = target_market_and_goods_overlap(
        existing_trademark["goods_services"], proposed_goods_services
    )

    # condition_1A_satisfieds = is_exact_match(existing_trademark['trademark_name'].strip().lower(), proposed_name.strip().lower())
    # st.write(f"Exact Match: {condition_1A_satisfieds}")

    # condition_1B_satisfieds = is_semantically_equivalents(existing_trademark['trademark_name'].strip().lower(), proposed_name.strip().lower())
    # st.write(f"Semantically equivalents : {condition_1B_satisfieds}")

    # condition_1C_satisfieds = is_phonetically_equivalents(existing_trademark['trademark_name'], proposed_name)
    # st.write(f"Phonetically equivalents : {condition_1C_satisfieds}")

    # condition_3_satisfieds = target_market_and_goods_overlaps(existing_trademark['goods_services'], proposed_goods_services)
    # st.write(f"Goods and services match's : {condition_3_satisfieds}")

    # Clean and standardize the trademark status
    status = existing_trademark["status"].strip().lower()

    # Check for 'Cancelled' or 'Abandoned' status
    if any(keyword in status for keyword in ["cancelled", "abandoned", "expired"]):
        conflict_grade = "Low"
        reasoning = "The existing trademark status is 'Cancelled' or 'Abandoned.'"
    else:
        points = sum(
            [
                condition_1_satisfied,  # 1 point if any Condition 1 is satisfied
                condition_2_satisfied,  # 1 point if Condition 2 is satisfied
                condition_3_satisfied,  # 1 point if Condition 3 is satisfied
            ]
        )

        # Determine conflict grade based on points
        if points == 3:
            conflict_grade = "High"
        elif points == 2:
            conflict_grade = "Moderate"
        elif points == 1:
            conflict_grade = "Low"
        else:
            conflict_grade = "None"

        if condition_1_satisfied:
            condition_1_details = []
            if condition_1A_satisfied:
                condition_1_details.append("Exact character-for-character match")
            if condition_1B_satisfied:
                condition_1_details.append("Semantically equivalent")
            if condition_1C_satisfied:
                condition_1_details.append("Phonetically equivalent")
            if condition_1D_satisfied:
                condition_1_details.append(
                    "First two or more words are phonetically equivalent"
                )
            if condition_1E_satisfied:
                condition_1_details.append(
                    "Proposed name is the first word of the existing trademark"
                )

        # Generate detailed reasoning for Condition 1
        if condition_1_satisfied:
            condition_1_reasoning = (
                f"Condition 1: Satisfied - {', '.join(condition_1_details)}."
            )
        else:
            condition_1_reasoning = "Condition 1: Not Satisfied."

        # Reasoning
        reasoning = (
            f"{condition_1_reasoning} \n"
            f"Condition 2: {'Satisfied' if condition_2_satisfied else 'Not Satisfied'} - Overlap in class numbers.\n"
            f"Condition 3: {'Satisfied' if condition_3_satisfied else 'Not Satisfied'} - Overlap in goods/services and target market."
        )

    if existing_trademark["design_phrase"] == "No Design phrase presented in document":
        design_label = "Word"
    else:
        design_label = "Design"

    if condition_1_satisfied and condition_2_satisfied and condition_3_satisfied:
        return {
            "Trademark Name , Class Number": f"{existing_trademark['trademark_name']} , {existing_trademark['international_class_number']}",
            "Trademark name": existing_trademark["trademark_name"],
            "Trademark Status": existing_trademark["status"],
            "Trademark Owner": existing_trademark["owner"],
            "Trademark class Number": existing_trademark["international_class_number"],
            "Trademark serial number": existing_trademark["serial_number"],
            "Serial / Registration Number": f"{existing_trademark['serial_number']} / {existing_trademark['registration_number']}",
            "Trademark registration number": existing_trademark["registration_number"],
            "Trademark design phrase": existing_trademark["design_phrase"],
            "Word/Design": design_label,
            "conflict_grade": conflict_grade,
            "reasoning": reasoning,
            "Mark": "   ✔️",
            "Class": "   ✔️",
            "Goods/Services": "   ✔️",
            "Direct Hit": " ",
        }
        
    elif condition_1_satisfied and condition_2_satisfied:
        return {
            "Trademark Name , Class Number": f"{existing_trademark['trademark_name']} , {existing_trademark['international_class_number']}",
            "Trademark name": existing_trademark["trademark_name"],
            "Trademark Status": existing_trademark["status"],
            "Trademark Owner": existing_trademark["owner"],
            "Trademark class Number": existing_trademark["international_class_number"],
            "Trademark serial number": existing_trademark["serial_number"],
            "Serial / Registration Number": f"{existing_trademark['serial_number']} / {existing_trademark['registration_number']}",
            "Trademark registration number": existing_trademark["registration_number"],
            "Trademark design phrase": existing_trademark["design_phrase"],
            "Word/Design": design_label,
            "conflict_grade": conflict_grade,
            "reasoning": reasoning,
            "Mark": "   ✔️",
            "Class": "   ✔️",
            "Goods/Services": "  ",
            "Direct Hit": " ",
        }
        
    elif condition_2_satisfied and condition_3_satisfied:
        return {
            "Trademark Name , Class Number": f"{existing_trademark['trademark_name']} , {existing_trademark['international_class_number']}",
            "Trademark name": existing_trademark["trademark_name"],
            "Trademark Status": existing_trademark["status"],
            "Trademark Owner": existing_trademark["owner"],
            "Trademark class Number": existing_trademark["international_class_number"],
            "Trademark serial number": existing_trademark["serial_number"],
            "Serial / Registration Number": f"{existing_trademark['serial_number']} / {existing_trademark['registration_number']}",
            "Trademark registration number": existing_trademark["registration_number"],
            "Trademark design phrase": existing_trademark["design_phrase"],
            "Word/Design": design_label,
            "conflict_grade": conflict_grade,
            "reasoning": reasoning,
            "Mark": " ",
            "Class": "   ✔️",
            "Goods/Services": "   ✔️",
            "Direct Hit": " ",
        }

    else:
        return {
            "Trademark Name , Class Number": f"{existing_trademark['trademark_name']} , {existing_trademark['international_class_number']}",
            "Trademark name": existing_trademark["trademark_name"],
            "Trademark Status": existing_trademark["status"],
            "Trademark Owner": existing_trademark["owner"],
            "Trademark class Number": existing_trademark["international_class_number"],
            "Trademark serial number": existing_trademark["serial_number"],
            "Serial / Registration Number": f"{existing_trademark['serial_number']} / {existing_trademark['registration_number']}",
            "Trademark registration number": existing_trademark["registration_number"],
            "Trademark design phrase": existing_trademark["design_phrase"],
            "Word/Design": design_label,
            "conflict_grade": conflict_grade,
            "reasoning": reasoning,
            "Mark": " ",
            "Class": "   ✔️",
            "Goods/Services": " ",
            "Direct Hit": " ",
        }

def replace_disallowed_words(text):
    disallowed_words = {
        "sexual": "xxxxxx",
        "sex": "xxx",
    }
    for word, replacement in disallowed_words.items():
        text = text.replace(word, replacement)
    # Ensure single paragraph output
    text = " ".join(text.split())
    return text


def assess_conflict(
    existing_trademark: List[Dict[str, Union[str, List[int]]]],
    proposed_name: str,
    proposed_class: str,
    proposed_goods_services: str,
) -> List[Dict[str, int]]:

    import phonetics
    from sentence_transformers import util
    from rapidfuzz import fuzz

    def normalize_text_name(text):
        """Normalize text by converting to lowercase, removing special characters, and standardizing whitespace."""
        # Remove punctuation except hyphens and spaces
        # text = re.sub(r"[^\w\s-’]", "", text)
        # Convert to lowercase
        text = re.sub(r"’", " ", text)
        text = text.lower()
        # Standardize whitespace
        return " ".join(text.split())

    # Clean and standardize the trademark status
    status = existing_trademark["status"].strip().lower()
    # Check for 'Cancelled' or 'Abandoned' status
    if any(keyword in status for keyword in ["cancelled", "abandoned", "expired"]):
        conflict_grade = "Low"
        reasoning = "The existing trademark status is 'Cancelled' or 'Abandoned.'"
    else:

        existing_trademark_name = normalize_text_name(
            existing_trademark["trademark_name"]
        )
        proposed_name = normalize_text_name(proposed_name)

        # Phonetic Comparison
        existing_phonetic = phonetics.metaphone(existing_trademark_name)
        proposed_phonetic = phonetics.metaphone(proposed_name)
        phonetic_match = existing_phonetic == proposed_phonetic

        # Semantic Similarity
        existing_embedding = semantic_model.encode(
            existing_trademark_name, convert_to_tensor=True
        )
        proposed_embedding = semantic_model.encode(
            proposed_name, convert_to_tensor=True
        )
        semantic_similarity = util.cos_sim(
            existing_embedding, proposed_embedding
        ).item()

        # String Similarity
        string_similarity = fuzz.ratio(existing_trademark_name, proposed_name)

        def is_substring_match(name1, name2):
            return name1.lower() in name2.lower() or name2.lower() in name1.lower()

        substring_match = is_substring_match(existing_trademark_name, proposed_name)

        def has_shared_word(name1, name2):
            words1 = set(name1.lower().split())
            words2 = set(name2.lower().split())
            return not words1.isdisjoint(words2)

        shared_word = has_shared_word(existing_trademark_name, proposed_name)

        from fuzzywuzzy import fuzz

        def is_phonetic_partial_match(name1, name2, threshold=55):
            return fuzz.partial_ratio(name1.lower(), name2.lower()) >= threshold

        phonetic_partial_match = is_phonetic_partial_match(
            existing_trademark_name, proposed_name
        )

        # st.write(f"Shared word : {existing_trademark_name} : {shared_word}")
        # st.write(f"Phonetic partial match : {existing_trademark_name} : {phonetic_partial_match}")
        # st.write(f"Substring match : {existing_trademark_name} : {substring_match}")

        # Decision Logic
        if (
            phonetic_match
            or substring_match
            or shared_word
            or semantic_similarity >= 0.5
            or string_similarity >= 55
            or phonetic_partial_match >= 55
        ):
            conflict_grade = "Name-Match"
        else:
            conflict_grade = "Low"

        semantic_similarity = semantic_similarity * 100

        # Reasoning
        reasoning = (
            f"Condition 1: {'Satisfied' if phonetic_match else 'Not Satisfied'} - Phonetic match found.\n"
            f"Condition 2: {'Satisfied' if substring_match else 'Not Satisfied'} - Substring match found.\n"
            f"Condition 3: {'Satisfied' if shared_word else 'Not Satisfied'} - Substring match found.\n"
            f"Condition 4: {'Satisfied' if phonetic_partial_match >= 55 else 'Not Satisfied'} - String similarity is ({round(phonetic_partial_match)}%).\n"
            f"Condition 5: {'Satisfied' if semantic_similarity >= 50 else 'Not Satisfied'} - Semantic similarity is ({round(semantic_similarity)}%).\n"
            f"Condition 6: {'Satisfied' if string_similarity >= 55 else 'Not Satisfied'} - String similarity is ({round(string_similarity)}%).\n"
        )

    if existing_trademark["design_phrase"] == "No Design phrase presented in document":
        design_label = "Word"
    else:
        design_label = "Design"

    return {
        "Trademark Name , Class Number": f"{existing_trademark['trademark_name']} , {existing_trademark['international_class_number']}",
        "Trademark name": existing_trademark["trademark_name"],
        "Trademark Status": existing_trademark["status"],
        "Trademark Owner": existing_trademark["owner"],
        "Trademark class Number": existing_trademark["international_class_number"],
        "Trademark serial number": existing_trademark["serial_number"],
        "Serial / Registration Number": f"{existing_trademark['serial_number']} / {existing_trademark['registration_number']}",
        "Trademark registration number": existing_trademark["registration_number"],
        "Trademark design phrase": existing_trademark["design_phrase"],
        "Word/Design": design_label,
        "conflict_grade": conflict_grade,
        "reasoning": reasoning,
        "Mark": " ",
        "Class": " ",
        "Goods/Services": " ",
        "Direct Hit": "   ✔️",

    }


import os
import json
from openai import AzureOpenAI


# Function to compare trademarks
def compare_trademarks2(
    existing_trademark: List[Dict[str, Union[str, List[int]]]],
    proposed_name: str,
    proposed_class: str,
    proposed_goods_services: str,
) -> List[Dict[str, Union[str, int]]]:
    proposed_classes = [int(c.strip()) for c in proposed_class.split(",")]

    # Prepare the messages for the Azure OpenAI API
    messages = [
        {
            "role": "system",
            "content": """  
            You are a trademark attorney tasked with determining a conflict grade based on the given conditions.  
            
            **Additional Instructions:**  
            
            - Consider if the proposed trademark name appears anywhere within the existing trademark name, or if significant parts of the existing trademark name appear in the proposed name.  
            - Evaluate shared words between trademarks, regardless of their position.  
            - Assess phonetic similarities, including partial matches.  
            - Consider the overall impression created by the trademarks, including similarities in appearance, sound, and meaning.  
            
            Follow the conflict grading criteria as previously outlined, assigning "Name-Match" or "Low" based on your analysis.  
            """,
        },
        {
            "role": "user",
            "content": f"""  
            Evaluate the potential conflict between the following existing trademarks and the proposed trademark.  
            
            **Proposed Trademark:**  
            - Name: "{proposed_name}"  
            
            **Existing Trademarks:**  
            - Name: "{existing_trademark['trademark_name']}"  
            - Status: "{existing_trademark['status']}"
            
            **Instructions:**  
            1. Review the proposed and existing trademark data.  
            2. Determine if the trademarks are likely to cause confusion based on the Trademark name such as Phonetic match, Semantic similarity and String similarity.  
            3. Return the output with Conflict Grade only as 'Name-Match' or 'Low', based on the reasoning. 
            4. Provide reasoning for each Conflict Grade.
            5. Special Case: If the existing trademark status is "Cancelled" or "Abandoned," it will automatically be considered as Conflict Grade: Low.  
            
            **Output Format:**  
                Existing Name: Name of the existing trademark.
                Reasoning: Reasoning for the conflict grade.
                Conflict Grade: Name-Match
        """,
        },
    ]

    # Initialize the Azure OpenAI client
    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    api_key = os.getenv("AZURE_API_KEY")

    if not azure_endpoint or not api_key:
        raise ValueError(
            "Azure endpoint or API key is not set in environment variables."
        )

    client = AzureOpenAI(
        azure_endpoint=azure_endpoint, api_key=api_key, api_version="2024-10-01-preview"
    )

    # Call Azure OpenAI to get the response
    try:
        response_reasoning = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            temperature=0,
            max_tokens=500,
            top_p=1,
        )

        # Extract the content from the response
        reasoning_content = response_reasoning.choices[0].message.content
        conflict_grade = reasoning_content.split("Conflict Grade:", 1)[1].strip()
        st.write(reasoning_content)

        return conflict_grade

    except Exception as e:
        print(f"Error while calling Azure OpenAI API: {e}")
        return []


def extract_proposed_trademark_details(
    file_path: str,
) -> Dict[str, Union[str, List[int]]]:
    """Extract proposed trademark details from the given input format"""
    proposed_details = {}
    with fitz.open(file_path) as pdf_document:
        if pdf_document.page_count > 0:
            page = pdf_document.load_page(0)
            page_text = preprocess_text(page.get_text())
            if "Mark Searched:" not in page_text:
                page = pdf_document.load_page(1)
                page_text = preprocess_text(page.get_text())

    name_match = re.search(
        r"Mark Searched:\s*(.*?)(?=\s*Client Name:)",
        page_text,
        re.IGNORECASE | re.DOTALL,
    )
    if name_match:
        proposed_details["proposed_trademark_name"] = name_match.group(1).strip()

    if "Goods/Services:" in page_text:
        goods_services_match = re.search(
            r"Goods/Services:\s*(.*?)(?=\s*Trademark Research Report)",
            page_text,
            re.IGNORECASE | re.DOTALL,
        )
    else:
        goods_services_match = re.search(
            r"Goods and Services:\s*(.*?)(?=\s*Order Info)",
            page_text,
            re.IGNORECASE | re.DOTALL,
        )

    if goods_services_match:
        proposed_details["proposed_goods_services"] = goods_services_match.group(
            1
        ).strip()

    # Use LLM to find the international class number based on goods & services
    if "proposed_goods_services" in proposed_details:
        goods_services = proposed_details["proposed_goods_services"]
        class_numbers = find_class_numbers(goods_services)
        proposed_details["proposed_nice_classes_number"] = class_numbers

    return proposed_details


def find_class_numbers(goods_services: str) -> List[int]:
    """Use LLM to find the international class numbers based on goods & services"""
    # Initialize AzureChatOpenAI

    from openai import AzureOpenAI

    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    api_key = os.getenv("AZURE_API_KEY")

    client = AzureOpenAI(
        azure_endpoint=azure_endpoint,
        api_key=api_key,
        api_version="2024-10-01-preview",
    )

    messages = [
        {
            "role": "system",
            "content": "You are a helpful assistant for finding the International class number of provided Goods & Services.",
        },
        {
            "role": "user",
            "content": "The goods/services are: IC 003: SKIN CARE PREPARATIONS; COSMETICS; BABY CARE PRODUCTS, NAMELY, SKIN SOAPS, BABY WASH, BABY BUBBLE BATH, BABY LOTIONS, BABY SHAMPOOS; SKIN CLEANSERS; BABY WIPES; NON− MEDICATED DIAPER RASH OINTMENTS AND LOTIONS; SKIN LOTIONS, CREAMS, MOISTURIZERS, AND OILS; BODY WASH; BODY SOAP; DEODORANTS; PERFUME; HAIR CARE PREPARATIONS. Find the international class numbers.",
        },
        {"role": "assistant", "content": "The international class numbers : 03"},
        {
            "role": "user",
            "content": "The goods/services are: LUGGAGE AND CARRYING BAGS; SUITCASES, TRUNKS, TRAVELLING BAGS, SLING BAGS FOR CARRYING INFANTS, SCHOOL BAGS; PURSES; WALLETS; RETAIL AND ONLINE RETAIL SERVICES. Find the international class numbers.",
        },
        {"role": "assistant", "content": "The international class numbers : 18,35"},
        {
            "role": "user",
            "content": "The goods/services are: CLASS 3: ANTIPERSPIRANTS AND DEODORANTS. (PLEASE INCLUDE CLASSES 5 AND 35 IN THE SEARCH SCOPE). Find the international class numbers.",
        },
        {"role": "assistant", "content": "The international class numbers : 03,05,35"},
        {
            "role": "user",
            "content": "The goods/services are: VITAMIN AND MINERAL SUPPLEMENTS. Find the international class numbers.",
        },
        {"role": "assistant", "content": "The international class numbers : 05"},
        {
            "role": "user",
            "content": f"The goods/services are: {goods_services}. Find the international class numbers.",
        },
    ]
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=messages,
        temperature=0.5,
        max_tokens=150,
    )

    class_numbers_str = response.choices[0].message.content

    # Extracting class numbers and removing duplicates
    class_numbers = re.findall(
        r"(?<!\d)\d{2}(?!\d)", class_numbers_str
    )  # Look for two-digit numbers
    class_numbers = ",".join(
        set(class_numbers)
    )  # Convert to set to remove duplicates, then join into a single string

    return class_numbers


def extract_proposed_trademark_details2(
    file_path: str,
) -> Dict[str, Union[str, List[int]]]:
    """Extract proposed trademark details from the first page of the document"""
    proposed_details = {}
    with fitz.open(file_path) as pdf_document:
        if pdf_document.page_count > 0:
            page = pdf_document.load_page(0)
            page_text = preprocess_text(page.get_text())

            name_match = re.search(r"Name:\s*(.*?)(?=\s*Nice Classes:)", page_text)
            if name_match:
                proposed_details["proposed_trademark_name"] = name_match.group(
                    1
                ).strip()

            nice_classes_match = re.search(
                r"Nice Classes:\s*(\d+(?:,\s*\d+)*)", page_text
            )
            if nice_classes_match:
                proposed_details["proposed_nice_classes_number"] = (
                    nice_classes_match.group(1).strip()
                )

            goods_services_match = re.search(
                r"Goods & Services:\s*(.*?)(?=\s*Registers|$)",
                page_text,
                re.IGNORECASE | re.DOTALL,
            )
            if goods_services_match:
                proposed_details["proposed_goods_services"] = (
                    goods_services_match.group(1).strip()
                )

    return proposed_details


def list_conversion(proposed_class: str) -> List[int]:

    from openai import AzureOpenAI

    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    api_key = os.getenv("AZURE_API_KEY")

    client = AzureOpenAI(
        azure_endpoint=azure_endpoint,
        api_key=api_key,
        api_version="2024-10-01-preview",
    )

    messages = [
        {
            "role": "system",
            "content": "You are a helpful assistant for converting the class number string into python list of numbers.\n Respond only with python list. Example : [18,35]",
        },
        {
            "role": "user",
            "content": "The class number are: 15,89. convert the string into python list of numbers.",
        },
        {"role": "assistant", "content": "[15,89]"},
        {
            "role": "user",
            "content": f"The class number are: {proposed_class}. convert the string into python list of numbers.",
        },
    ]
    # messages = [
    # {
    #     "role": "system",
    #     "content": "You are a helpful assistant that converts strings of class numbers into Python lists of integers."
    # },
    # {
    # "role": "user",
    # "content": f"""
    #     Convert the following string of class numbers into a Python list of integers.

    #     **Instructions:**

    #     - The input is a string of numbers separated by commas (e.g., `15,89`).
    #     - **Respond only** with a Python list of integers (e.g., `[15, 89]`).
    #     - Do not include any additional text or commentary.
    #     - Ensure the numbers are integers, not strings.

    #     **Example:**

    #     - Input: "15,89"
    #     - Response: [15, 89]

    #     **Input:**

    #     "{proposed_class}"
    #     """
    # }
    # ]

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=messages,
        temperature=0,
        max_tokens=150,
    )

    lst_class = response.choices[0].message.content
    class_value = ast.literal_eval(lst_class)

    return class_value

# TAMIL CODE START'S HERE-------------------------------------------------------------------------------------------------------------------------

import os
from openai import AzureOpenAI
import json
import re

def get_azure_client():
    """Initialize and return the Azure OpenAI client."""
    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    api_key = os.getenv("AZURE_API_KEY")
    
    client = AzureOpenAI(
        azure_endpoint=azure_endpoint,
        api_key=api_key,
        api_version="2024-10-01-preview",
    )
    return client

def validate_trademark_relevance(conflicts_array, proposed_goods_services):
    """
    Pre-filter trademarks that don't have similar or identical goods/services
    This function is implemented in code rather than relying on GPT
    
    Args:
        conflicts_array: List of trademark conflicts
        proposed_goods_services: Goods/services of the proposed trademark
        
    Returns:
        filtered_conflicts: List of relevant trademark conflicts
        excluded_count: Number of trademarks excluded
    """
    # Parse conflicts_array if it's a string (assuming JSON format)
    if isinstance(conflicts_array, str):
        try:
            conflicts = json.loads(conflicts_array)
        except json.JSONDecodeError:
            # If it's not valid JSON, try to parse it as a list of dictionaries
            conflicts = eval(conflicts_array) if conflicts_array.strip().startswith("[") else []
    else:
        conflicts = conflicts_array
    
    # Initialize lists for relevant and excluded trademarks
    relevant_conflicts = []
    excluded_count = 0
    
    # Define a function to check similarity between goods/services
    def is_similar_goods_services(existing_goods, proposed_goods):
        # Convert to lowercase for case-insensitive comparison
        existing_lower = existing_goods.lower()
        proposed_lower = proposed_goods.lower()
        
        # Check for exact match
        if existing_lower == proposed_lower:
            return True
        
        # Check if one contains the other
        if existing_lower in proposed_lower or proposed_lower in existing_lower:
            return True
        
        # Check for overlapping keywords
        # Extract significant keywords from both descriptions
        existing_keywords = set(re.findall(r'\b\w+\b', existing_lower))
        proposed_keywords = set(re.findall(r'\b\w+\b', proposed_lower))
        
        # Remove common stop words
        stop_words = {'and', 'or', 'the', 'a', 'an', 'in', 'on', 'for', 'of', 'to', 'with'}
        existing_keywords = existing_keywords - stop_words
        proposed_keywords = proposed_keywords - stop_words
        
        # Calculate keyword overlap
        if len(existing_keywords) > 0 and len(proposed_keywords) > 0:
            overlap = len(existing_keywords.intersection(proposed_keywords))
            overlap_ratio = overlap / min(len(existing_keywords), len(proposed_keywords))
            
            # If significant overlap (more than 30%), consider them similar
            if overlap_ratio > 0.3:
                return True
        
        return False
    
    # Process each conflict
    for conflict in conflicts:
        # Ensure conflict has goods/services field
        if 'goods_services' in conflict:
            if is_similar_goods_services(conflict['goods_services'], proposed_goods_services):
                relevant_conflicts.append(conflict)
            else:
                excluded_count += 1
        else:
            # If no goods/services field, include it for safety
            relevant_conflicts.append(conflict)
    
    return relevant_conflicts, excluded_count

def filter_by_gpt_response(conflicts, gpt_json):
    """
    Removes trademarks that GPT flagged as lacking goods/services overlap.
    
    Args:
        conflicts: Original list of trademark conflicts
        gpt_json: JSON object from GPT with 'results' key
    
    Returns:
        Filtered list of conflicts that GPT identified as overlapping
    """
    # Parse the GPT response if it's a string
    if isinstance(gpt_json, str):
        try:
            gpt_json = json.loads(gpt_json)
        except json.JSONDecodeError:
            # If JSON is invalid, keep original conflicts
            return conflicts
    
    gpt_results = gpt_json.get("results", [])
    
    # Build a set of marks with overlap for quick membership checking
    overlapping_marks = {
        result["mark"]
        for result in gpt_results
        if result.get("overlap") is True
    }
    
    # Retain conflicts only if they appear in overlapping_marks
    filtered_conflicts = [
        c for c in conflicts
        if c.get("mark") in overlapping_marks
    ]
    
    return filtered_conflicts

# def initial_mark_analysis(conflicts_array, proposed_name, proposed_class, proposed_goods_services):
#     """
#     Perform Steps 1-6: Initial Mark Analysis
#     - First filter out irrelevant trademarks
#     - Then send only relevant trademarks to GPT for analysis
#     """
  
#     relevant_conflicts, excluded_count = validate_trademark_relevance(conflicts_array, proposed_goods_services)
    
   
#     system_prompt = """
#     You are a trademark expert attorney specializing in trademark opinion writing. Analyze the provided trademark data and provide a professional opinion on registration and use risks.
    
#     Follow these steps for your analysis:
    
#     Step 1: Verify and Deconstruct the Compound Mark
#     - Confirm if the proposed trademark is a compound mark (combination of words/elements).
#     - Deconstruct it into its formative components.
#     - Example: "MOUNTAIN FRESH" → "MOUNTAIN" and "FRESH"
    
#     Step 2: Identify Identical Trademarks
#     - List existing trademarks with identical names to the proposed trademark.
#     - Only consider trademarks with identical or similar goods/services.
    
#     Step 3: Identify Phonetically/Semantically Equivalent Marks
#     - List marks that sound similar or have similar meanings to the proposed trademark.
#     - Only consider trademarks with identical or similar goods/services.
    
#     Step 4: Identify Marks with One-Letter Differences
#     - List similar marks that differ by one letter from the proposed trademark.
#     - Only consider trademarks with identical or similar goods/services.
    
#     Step 5: Identify Marks with Two-Letter Differences
#     - List similar marks that differ by two letters from the proposed trademark.
#     - Only consider trademarks with identical or similar goods/services.
    
#     Step 6: Perform Crowded Field Analysis
#     - If Steps 4 and 5 yield more than 20 marks, check their ownership.
#     - Calculate the percentage of marks with different owners.
#     - If more than 50% have different owners, consider it a crowded field.
#     - If it's a crowded field, the final risk assessment should be reduced by one level.
    
#     IMPORTANT: 
#     - We have already filtered out trademarks with unrelated goods/services. 
#     - All trademarks in your input ARE relevant to the proposed trademark's goods/services.
#     - Focus on matches that are closely related to the entire trademark name, not just individual components.
    
#     YOUR RESPONSE MUST END WITH A JSON SUMMARY in this exact format:
#     {
#       "results": [
#         {
#           "mark": "[TRADEMARK NAME]",
#           "owner": "[OWNER NAME]",
#           "goods_services": "[GOODS/SERVICES DESCRIPTION]",
#           "overlap": true,
#           "risk_level": "[HIGH|MEDIUM|LOW]",
#           "class_match": true|false,
#           "goods_services_match": true|false
#         },
#         ...additional marks...
#       ],
#       "summary": {
#         "identical_count": [NUMBER],
#         "phonetic_count": [NUMBER],
#         "one_letter_count": [NUMBER],
#         "two_letter_count": [NUMBER],
#         "crowded_field": {
#           "is_crowded": true|false,
#           "percentage": [PERCENTAGE],
#           "explanation": "[EXPLANATION]"
#         }
#       }
#     }
#     """
    
#     client = get_azure_client()
    
   
#     user_message = f"""
#     Trademark Details:
#     {json.dumps(relevant_conflicts, indent=2)}
    
#     Analyze the Proposed Trademark "{proposed_name}" focusing on Steps 1-6: Initial mark analysis.
#     Mark Searched = {proposed_name}
#     Classes Searched = {proposed_class}
#     Goods and Services = {proposed_goods_services}
    
#     Note: {excluded_count} trademarks with unrelated goods/services have already been filtered out.
    
#     For each mark, determine:
#     - "Class Match" (True/False): Whether the mark's class exactly matches the proposed class "{proposed_class}".
#     - "Goods & Services Match" (True/False): Whether the mark's goods/services are similar to the proposed goods/services "{proposed_goods_services}".
    
#     REMEMBER: 
#     - End your response with the JSON summary as specified in the instructions.
#     - Include owner names and goods/services details for each mark.
#     - Focus on matches to the entire trademark name, not just components.
#     """
    
#     try:
#         response = client.chat.completions.create(
#             model="gpt-4o",
#             messages=[
#                 {"role": "system", "content": system_prompt},
#                 {"role": "user", "content": user_message}
#             ],
#             temperature=0.0,
#         )
        
    
#         if response.choices and len(response.choices) > 0:
#             content = response.choices[0].message.content
            
           
#             json_match = re.search(r'```json\s*({[\s\S]*?})\s*```|({[\s\S]*?"summary"\s*:[\s\S]*?})', content)
#             if json_match:
#                 json_str = json_match.group(1) or json_match.group(2)
#                 try:
#                     json_data = json.loads(json_str)
                   
#                     return {
#                         "analysis": content,
#                         "json_data": json_data
#                     }
#                 except json.JSONDecodeError:
#                     pass
            
#             return content
#         else:
#             return "Error: No response received from the language model."
#     except Exception as e:
#         return f"Error during initial mark analysis: {str(e)}"

# def component_formative_mark_analysis(conflicts_array, proposed_name, proposed_class, proposed_goods_services):
#     """
#     Perform Step 8: Component (Formative) Mark Analysis
#     - Pre-filter trademarks before sending to GPT
#     """
    
#     relevant_conflicts, excluded_count = validate_trademark_relevance(conflicts_array, proposed_goods_services)
    
  
#     system_prompt = """
# You are a trademark expert attorney specializing in trademark opinion writing.

# Perform Step 8: Component (Formative) Mark Analysis using the following structure:

# Step 8.a: Identify and Deconstruct the Compound Mark
# - Confirm if the proposed trademark is a compound mark (combination of words/elements).
# - Deconstruct it into its formative components.
# - Example: For "POWERHOLD," identify the components "POWER" and "HOLD".

# FOR EACH FORMATIVE COMPONENT, perform the following detailed analysis:

# Step 8.b: Identical Marks Analysis for Each Component
# - Only list trademarks that are identical to each individual formative component AND cover identical or similar goods/services.
# - Example: For "POWERHOLD," analyze "POWER" trademarks and "HOLD" trademarks separately.
# - If no identical marks pass validation for a component, state: "No identical trademarks covering similar goods/services were identified for [COMPONENT]."

# Step 8.c: Phonetic and Semantic Equivalents for Each Component
# - Only list trademarks that are phonetically or semantically similar to each formative component AND cover identical or similar goods/services.
# - Example: For "POWER," phonetically similar marks might include "POWR," "POWUR," or "PAWER." 
# - Evaluate whether these similar marks overlap in goods/services and assess the likelihood of confusion.

# Step 8.d: Marks with Letter Differences for Each Component
# Step 8.d.1: One-Letter Differences
# - Only list trademarks that differ by one letter from each formative component AND cover identical or similar goods/services.
# - Example: For "POWER," consider marks like "POWIR" or "POSER."
# - Assess the impact of these differences on consumer perception and the likelihood of confusion.

# Step 8.d.2: Two-Letter Differences
# - List ONLY trademarks that differ by two letters from each formative component AND cover relevant goods/services.
# - Example: For "POWER," consider "POWTR" or "PIWER."
# - Evaluate whether these differences create confusion in meaning or pronunciation.

# Step 8.e: Component Distinctiveness Analysis
# - For each component, classify its distinctiveness as Generic, Descriptive, Suggestive, Arbitrary, or Fanciful.
# - Consider the component in relation to the specific goods/services.
# - Example: For "POWER" in electrical equipment, it would be descriptive; for food services, it would be arbitrary.

# Step 8.f: Functional/Conceptual Relationship Analysis
# - For compound marks, analyze how the meaning of one component might relate functionally to another component in EXISTING marks.
# - Example: For "MIRAGRIP," identify marks where a component has a functional relationship similar to how "MIRA" relates to "GRIP" (e.g., "VISIONHOLD," "WONDERCLUTCH").
# - Only include marks with relevant goods/services.
# - Document the functional relationship between components and why they create similar commercial impressions.

# IMPORTANT: 
# - We have already filtered out ALL trademarks with unrelated goods/services. 
# - Your analysis should ONLY include trademarks with goods/services relevant to the proposed trademark.
# - Include owner names and goods/services details for each mark.

# YOUR RESPONSE MUST END WITH A JSON SUMMARY in this exact format:
# {
#   "components": [
#     {
#       "component": "[COMPONENT NAME]",
#       "results": [
#         {
#           "mark": "[TRADEMARK NAME]",
#           "owner": "[OWNER NAME]",
#           "goods_services": "[GOODS/SERVICES DESCRIPTION]",
#           "overlap": true
#         },
#         ...additional marks for this component...
#       ],
#       "distinctiveness": "[GENERIC|DESCRIPTIVE|SUGGESTIVE|ARBITRARY|FANCIFUL]"
#     },
#     ...additional components...
#   ],
#   "crowded_field": {
#     "is_crowded": true|false,
#     "percentage": [PERCENTAGE],
#     "explanation": "[EXPLANATION]"
#   }
# }
# """
    
#     client = get_azure_client()
    
 
#     user_message = f"""
#     Trademark Details:
#     {json.dumps(relevant_conflicts, indent=2)}
    
#     Analyze the Proposed Trademark "{proposed_name}" focusing on Step 8: Component (Formative) Mark Analysis.
#     Mark Searched = {proposed_name}
#     Classes Searched = {proposed_class}
#     Goods and Services = {proposed_goods_services}
    
#     Note: {excluded_count} trademarks with unrelated goods/services have already been filtered out.
    
#     REMEMBER: 
#     - End your response with the JSON summary as specified in the instructions.
#     - Include owner names and goods/services details for each mark.
#     - Focus on analyzing each component separately.
#     """
    
#     try:
#         response = client.chat.completions.create(
#             model="gpt-4o",
#             messages=[
#                 {"role": "system", "content": system_prompt},
#                 {"role": "user", "content": user_message}
#             ],
#             temperature=0.0,
#         )
        
 
#         if response.choices and len(response.choices) > 0:
#             content = response.choices[0].message.content
            
       
#             json_match = re.search(r'```json\s*({[\s\S]*?})\s*```|({[\s\S]*?"components"\s*:[\s\S]*?})', content)
#             if json_match:
#                 json_str = json_match.group(1) or json_match.group(2)
#                 try:
#                     json_data = json.loads(json_str)
                  
#                     return {
#                         "analysis": content,
#                         "json_data": json_data
#                     }
#                 except json.JSONDecodeError:
#                     pass
            
#             return content
#         else:
#             return "Error: No response received from the language model."
#     except Exception as e:
#         return f"Error during component formative mark analysis: {str(e)}"

# def final_validation_and_assessment(conflicts_array, proposed_name, proposed_class, proposed_goods_services, step7_results, step8_results, excluded_count):
#     """
#     Perform Steps 9-11: Final Validation, Overall Risk Assessment, and Summary of Findings
#     - Pass the excluded_count to inform GPT about pre-filtering
#     """

#     system_prompt = """
#     You are a trademark expert attorney specializing in trademark opinion writing.
    
#     Perform Steps 9-11: Final Validation, Overall Risk Assessment, and Summary of Findings using the following structure:
    
#     Step 9: Final Validation Check
#     - All trademarks with unrelated goods/services have already been filtered out. No further filtering is needed.
    
#     Step 10: Overall Risk Assessment
#     - Integrate all findings from previous steps (Steps 1-8) to provide a single, comprehensive risk assessment.
#     - Assess the trademark's overall viability and risk on this scale:
#       * Low: Very few/no conflicts, highly distinctive mark
#       * Medium-Low: Some minor conflicts, moderately distinctive mark
#       * Medium: Several potential conflicts, average distinctiveness
#       * Medium-High: Numerous conflicts, limited distinctiveness
#       * High: Significant conflicts, minimal distinctiveness
#     - Consider these factors:
#       * Number and similarity of identical marks
#       * Number and similarity of phonetically/semantically equivalent marks
#       * Presence of marks with one or two-letter differences
#       * Crowded field status (if applicable, reduce risk by one level)
#       * Evidence of aggressive enforcement by owners of similar marks
#       * Distinctiveness of the compound mark and its components
#     - Focus the discussion on how the crowded field analysis contributed to risk reduction.
    
#     Step 11: Summary of Findings
#     - Summarize the overall trademark analysis, including:
#       * Likelihood of Conflicts
#       * Crowded Field Status (with numerical percentages)
#       * Distinctiveness Assessment
#     - Do NOT include recommendations in the summary.
    
#     YOUR RESPONSE MUST END WITH A JSON SUMMARY in this exact format:
#     {
#       "final_assessment": {
#         "overall_risk_level": "[HIGH|MEDIUM-HIGH|MEDIUM|MEDIUM-LOW|LOW]",
#         "crowded_field": {
#           "is_crowded": true|false,
#           "percentage": [PERCENTAGE],
#           "explanation": "[EXPLANATION]"
#         },
#         "identical_mark_count": [NUMBER],
#         "similar_mark_count": [NUMBER],
#         "key_conflicts": ["[TRADEMARK1]", "[TRADEMARK2]", ...]
#       }
#     }
#     """
    
#     client = get_azure_client()
    

#     step7_json = step7_results.get("json_data", {}) if isinstance(step7_results, dict) else {}
#     step8_json = step8_results.get("json_data", {}) if isinstance(step8_results, dict) else {}
    

#     step7_analysis = step7_results.get("analysis", step7_results) if isinstance(step7_results, dict) else step7_results
#     step8_analysis = step8_results.get("analysis", step8_results) if isinstance(step8_results, dict) else step8_results
    
  
#     user_message = f"""
#     Trademark Details:
#     - Proposed Trademark: {proposed_name}
#     - Classes Searched: {proposed_class}
#     - Goods and Services: {proposed_goods_services}
    
#     Previous Analysis Results:
    
#     --- Step 7 Results ---
#     {step7_analysis}
    
#     --- Step 8 Results ---
#     {step8_analysis}
    
#     Please complete the trademark analysis by performing Steps 9-11: Final Validation Check, Overall Risk Assessment, and Summary of Findings.
    
#     Note: {excluded_count} trademarks with unrelated goods/services were excluded from this analysis through pre-filtering.
    
#     REMEMBER: 
#     - End your response with the JSON summary as specified in the instructions.
#     - Focus the risk discussion on crowded field analysis.
#     - Include numerical percentages for crowded field analysis.
#     - Do NOT include recommendations.
#     """
    
#     try:
#         response = client.chat.completions.create(
#             model="gpt-4o",
#             messages=[
#                 {"role": "system", "content": system_prompt},
#                 {"role": "user", "content": user_message}
#             ],
#             temperature=0.0,
#         )
        
      
#         if response.choices and len(response.choices) > 0:
#             content = response.choices[0].message.content
            
         
#             json_match = re.search(r'```json\s*({[\s\S]*?})\s*```|({[\s\S]*?"final_assessment"\s*:[\s\S]*?})', content)
#             if json_match:
#                 json_str = json_match.group(1) or json_match.group(2)
#                 try:
#                     json_data = json.loads(json_str)
                   
#                     return {
#                         "analysis": content,
#                         "json_data": json_data
#                     }
#                 except json.JSONDecodeError:
#                     pass
            
#             return content
#         else:
#             return "Error: No response received from the language model."
#     except Exception as e:
#         return f"Error during final validation and assessment: {str(e)}"

def clean_and_format_opinion(comprehensive_opinion, json_data=None):
    """
    Process the comprehensive trademark opinion to:
    1. Maintain comprehensive listing of all relevant trademark hits
    2. Remove duplicated content while preserving all unique trademark references
    3. Format the opinion for better readability
    4. Ensure consistent structure with clear sections
    
    Args:
        comprehensive_opinion: Raw comprehensive opinion from previous steps
        json_data: Optional structured JSON data from previous steps
        
    Returns:
        A cleaned, formatted, and optimized trademark opinion
    """
    client = get_azure_client()
    
    system_prompt = """
    You are a trademark attorney specializing in clear, comprehensive trademark opinions.
    
    FORMAT THE TRADEMARK OPINION USING THE EXACT STRUCTURE PROVIDED BELOW:
    
    ```
REFINED TRADEMARK OPINION: [MARK NAME]
Class: [Class Number]
Goods and Services: [Goods/Services Description]

Section I: Comprehensive Trademark Hit Analysis
(a) Identical Marks:
| Trademark | Owner | Goods & Services | Status | Class | Class Match | Goods & Services Match |
|------------|--------|------------------|--------|------|-------------|------------------------|
| [Mark 1] | [Owner] | [Goods/Services] | [Status] | [Class] | [True/False] | [True/False] |

(b) One Letter and Two Letter Analysis:
| Trademark | Owner | Goods & Services | Status | Class | Difference Type | Class Match | Goods & Services Match |
|------------|--------|------------------|--------|------|----------------|-------------|------------------------|
| [Mark 1] | [Owner] | [Goods/Services] | [Status] | [Class] | [One/Two Letter] | [True/False] | [True/False] |

(c) Phonetically, Semantically & Functionally Similar Analysis:
| Trademark | Owner | Goods & Services | Status | Class | Similarity Type | Class Match | Goods & Services Match |
|------------|--------|------------------|--------|------|-----------------|-------------|------------------------|
| [Mark 1] | [Owner] | [Goods/Services] | [Status] | [Class] | [Phonetic/Semantic/Functional] | [True/False] | [True/False] |

Section II: Component Analysis
(a) Component Analysis:

Component 1: [First Component]
| Trademark | Owner | Goods & Services | Status | Class | Class Match | Goods & Services Match |
|-----------|--------|------------------|--------|-------|-------------|------------------------|
| [Mark 1] | [Owner] | [Goods/Services] | [Status] | [Class] | [True/False] | [True/False] |

Component A: [Second Component]
| Trademark | Owner | Goods & Services | Status | Class | Class Match | Goods & Services Match |
|-----------|--------|------------------|--------|-------|-------------|------------------------|
| [Mark 1] | [Owner] | [Goods/Services] | [Status] | [Class] | [True/False] | [True/False] |

(b) Crowded Field Analysis:
- **Total compound mark hits found**: [NUMBER]
- **Marks with different owners**: [NUMBER] ([PERCENTAGE]%)
- **Crowded Field Status**: [YES/NO]
- **Analysis**: 
  [DETAILED EXPLANATION OF FINDINGS INCLUDING RISK IMPLICATIONS IF FIELD IS CROWDED]

Section III: Risk Assessment and Summary

Descriptiveness:
- [KEY POINT ABOUT DESCRIPTIVENESS]

Aggressive Enforcement and Litigious Behavior:
- **Known Aggressive Owners**:
  * [Owner 1]: [Enforcement patterns]
  * [Owner 2]: [Enforcement patterns]
- **Enforcement Landscape**:
  * [KEY POINT ABOUT ENFORCEMENT LANDSCAPE]
  * [ADDITIONAL POINT ABOUT ENFORCEMENT LANDSCAPE]

Risk Category for Registration:
- **[REGISTRATION RISK LEVEL: HIGH/MEDIUM-HIGH/MEDIUM/MEDIUM-LOW/LOW]**
- [EXPLANATION OF REGISTRATION RISK LEVEL WITH FOCUS ON CROWDED FIELD ANALYSIS]

Risk Category for Use:
- **[USE RISK LEVEL: HIGH/MEDIUM-HIGH/MEDIUM/MEDIUM-LOW/LOW]**
- [EXPLANATION OF USE RISK LEVEL]
    ```

    **IMPORTANT INSTRUCTIONS:**
    1. Maintain ALL unique trademark references from the original opinion.
    2. Present trademarks in clear, easy-to-read tables following the format above.
    3. Ensure ALL findings from the original opinion are preserved but avoid redundancy.
    4. Include owner names and goods/services details for each mark.
    5. Include trademark search exclusions in the summary section.
    6. Ensure the final opinion is comprehensive yet concise.
    7. For each section, include all relevant trademarks without omission.
    8. Maintain the exact structure provided above with clear section headings.
    9. For each mark, determine and include:
       - "Class Match" (True/False): Whether the mark's class exactly matches the proposed trademark's class OR is in a coordinated/related class group.
       - "Goods & Services Match" (True/False): Whether the mark's goods/services are similar to the proposed trademark's goods/services.
    10. Follow the specified structure exactly:
        - Section I focuses on overall hits, including One/Two Letter Analysis
        - Section II focuses only on component hits
        - In Section II, perform Crowded Field Analysis focusing on owner diversity
    11. State "None" when no results are found for a particular subsection
    12. Do NOT include recommendations in the summary
    13. Include aggressive enforcement analysis in Section III with details on any owners known for litigious behavior
    14. IMPORTANT: When assessing "Class Match", consider not only exact class matches but also coordinated or related classes based on the goods/services.
    15. NEVER replace full goods/services descriptions with just class numbers in the output tables. Always include the complete goods/services text.
    """
    
    # Send the original opinion to be reformatted
    user_message = f"""
    Please reformat the following comprehensive trademark opinion according to the refined structure:
    
    Proposed Trademark: {json_data.get('proposed_name', 'N/A')}
    Class: {json_data.get('proposed_class', 'N/A')}
    Goods and Services: {json_data.get('proposed_goods_services', 'N/A')}
    
    Original Opinion:
    {comprehensive_opinion}
    
    Follow the exact structure provided in the instructions, ensuring all trademark references are maintained.
    
    For each mark in the tables, you must evaluate and include:
    1. Owner name
    2. Goods & Services description - ALWAYS include the FULL goods/services text, not just class numbers
    3. Class Match (True/False): 
       - Mark True if the mark's class exactly matches the proposed class "{json_data.get('proposed_class', 'N/A')}"
       - ALSO mark True if the mark's class is in a coordinated or related class grouping with the proposed class
       - First identify all coordinated classes based on the proposed goods/services: "{json_data.get('proposed_goods_services', 'N/A')}"
       - Then mark True for any mark in those coordinated classes
    4. Goods & Services Match (True/False): Compare the mark's goods/services to the proposed goods/services "{json_data.get('proposed_goods_services', 'N/A')}" and mark True if they are semantically similar.
    
    IMPORTANT REMINDERS FOR CROWDED FIELD ANALYSIS:
    - Include exact counts and percentages for:
      * Total compound mark hits found
      * Number and percentage of marks with different owners
      * Crowded Field Status (YES if >50% have different owners)
    - Clearly explain risk implications if field is crowded
    - Section I should include ALL hits (overall hits), not just compound mark hits
    - Section II should focus ONLY on compound mark hits
    - One and Two Letter Analysis should ONLY be in Section I, not Section II
    - If no results are found for a particular subsection, state "None"
    - Do NOT include recommendations in the summary
    - Include aggressive enforcement analysis in Section III with details on any owners known for litigious behavior
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            temperature=0.0,
        )
        
        # Extract and return the formatted opinion
        if response.choices and len(response.choices) > 0:
            formatted_opinion = response.choices[0].message.content
            
            # Filter out rows where both "Class Match" and "Goods & Services Match" are False
            filtered_opinion = []
            for line in formatted_opinion.splitlines():
                if "|" in line:  # Check if the line is part of a table
                    parts = line.split("|")
                    if len(parts) >= 7:  # Ensure the line has enough columns
                        # Check if this is a header row by looking for specific column header text
                        if "Class Match" in line or "Trademark" in line:
                            filtered_opinion.append(line)
                        else:
                            # For data rows, check the Class Match and Goods & Services Match values
                            class_match_idx = -3  # Second to last column
                            goods_services_match_idx = -1  # Last column
                            
                            class_match = "true" in parts[class_match_idx].strip().lower()
                            goods_services_match = "true" in parts[goods_services_match_idx].strip().lower()
                            
                            if class_match or goods_services_match:
                                filtered_opinion.append(line)
                    else:
                        # Include table formatting lines and other table parts
                        filtered_opinion.append(line)
                else:
                    # Include all non-table lines
                    filtered_opinion.append(line)

            # Join the filtered lines back into a single string
            filtered_opinion = "\n".join(filtered_opinion)
            
            return filtered_opinion
        else:
            return "Error: No response received from the language model."
    except Exception as e:
        return f"Error during opinion formatting: {str(e)}"
        
  
def consistency_check(mark, results):
    """
    Consistency checking function to ensure accuracy of analysis results.
    
    Args:
        mark: The proposed trademark name
        results: Raw analysis results
        
    Returns:
        Corrected and validated results
    """
    corrected_results = results.copy()
    
    # Ensure all entries have required fields
    required_fields = ['mark', 'owner', 'goods_services', 'status', 'class', 'class_match', 'goods_services_match']
    
    # Check identical marks
    for i, item in enumerate(corrected_results.get('identical_marks', [])):
        # Validate that mark name is indeed identical
        if item.get('mark', '').lower() != mark.lower():
            # Remove from identical marks if not actually identical
            corrected_results['identical_marks'][i] = None
        
        # Ensure all required fields exist
        for field in required_fields:
            if field not in item:
                if field == 'class_match' or field == 'goods_services_match':
                    corrected_results['identical_marks'][i][field] = False
                else:
                    corrected_results['identical_marks'][i][field] = "Unknown"
    
    # Remove None entries
    corrected_results['identical_marks'] = [item for item in corrected_results.get('identical_marks', []) if item is not None]
    
    # Similar checks for one_letter_marks
    for i, item in enumerate(corrected_results.get('one_letter_marks', [])):
        # Validate actual one letter difference
        if not is_one_letter_difference(item.get('mark', ''), mark):
            corrected_results['one_letter_marks'][i] = None
            
        # Ensure all required fields exist
        for field in required_fields:
            if field not in item and field != 'difference_type':
                if field == 'class_match' or field == 'goods_services_match':
                    corrected_results['one_letter_marks'][i][field] = False
                else:
                    corrected_results['one_letter_marks'][i][field] = "Unknown"
    
    # Remove None entries
    corrected_results['one_letter_marks'] = [item for item in corrected_results.get('one_letter_marks', []) if item is not None]
    
    # Similar checks for two_letter_marks
    for i, item in enumerate(corrected_results.get('two_letter_marks', [])):
        # Validate actual two letter difference
        if not is_two_letter_difference(item.get('mark', ''), mark):
            corrected_results['two_letter_marks'][i] = None
            
        # Ensure all required fields exist
        for field in required_fields:
            if field not in item and field != 'difference_type':
                if field == 'class_match' or field == 'goods_services_match':
                    corrected_results['two_letter_marks'][i][field] = False
                else:
                    corrected_results['two_letter_marks'][i][field] = "Unknown"
    
    # Remove None entries
    corrected_results['two_letter_marks'] = [item for item in corrected_results.get('two_letter_marks', []) if item is not None]
    
    # Check similar_marks
    for i, item in enumerate(corrected_results.get('similar_marks', [])):
        # Ensure all required fields exist
        for field in required_fields:
            if field not in item and field != 'similarity_type':
                if field == 'class_match' or field == 'goods_services_match':
                    corrected_results['similar_marks'][i][field] = False
                else:
                    corrected_results['similar_marks'][i][field] = "Unknown"
    
    return corrected_results


def is_one_letter_difference(mark1, mark2):
    """
    Check if two marks have a one letter difference.
    
    Args:
        mark1: First mark
        mark2: Second mark
        
    Returns:
        Boolean indicating if there's a one letter difference
    """
    # Handle case-insensitivity
    mark1 = mark1.lower()
    mark2 = mark2.lower()
    
    # If length difference is greater than 1, definitely not a one letter difference
    if abs(len(mark1) - len(mark2)) > 1:
        return False
    
    # Count differences
    differences = 0
    
    # Same length - check for substitution
    if len(mark1) == len(mark2):
        for c1, c2 in zip(mark1, mark2):
            if c1 != c2:
                differences += 1
                if differences > 1:
                    return False
    # Different length - check for insertion/deletion
    else:
        # Make sure mark1 is the shorter one for simplicity
        if len(mark1) > len(mark2):
            mark1, mark2 = mark2, mark1
            
        i, j = 0, 0
        while i < len(mark1) and j < len(mark2):
            if mark1[i] != mark2[j]:
                # Skip this character in the longer string
                j += 1
                differences += 1
                if differences > 1:
                    return False
            else:
                i += 1
                j += 1
                
    return differences == 1


def is_two_letter_difference(mark1, mark2):
    """
    Check if two marks have a two letter difference.
    
    Args:
        mark1: First mark
        mark2: Second mark
        
    Returns:
        Boolean indicating if there's a two letter difference
    """
    # Handle case-insensitivity
    mark1 = mark1.lower()
    mark2 = mark2.lower()
    
    # If length difference is greater than 2, definitely not a two letter difference
    if abs(len(mark1) - len(mark2)) > 2:
        return False
    
    # Use Levenshtein distance for accurate measurement
    return levenshtein_distance(mark1, mark2) == 2


def levenshtein_distance(s1, s2):
    """
    Calculate the Levenshtein distance between two strings.
    This measures the minimum number of single-character edits needed to change one string into another.
    
    Args:
        s1: First string
        s2: Second string
        
    Returns:
        The edit distance between the strings
    """
    if len(s1) < len(s2):
        return levenshtein_distance(s2, s1)

    if len(s2) == 0:
        return len(s1)

    previous_row = range(len(s2) + 1)
    for i, c1 in enumerate(s1):
        current_row = [i + 1]
        for j, c2 in enumerate(s2):
            insertions = previous_row[j + 1] + 1
            deletions = current_row[j] + 1
            substitutions = previous_row[j] + (c1 != c2)
            current_row.append(min(insertions, deletions, substitutions))
        previous_row = current_row
    
    return previous_row[-1]


def section_one_analysis(mark, class_number, goods_services, relevant_conflicts):  
    """
    Perform Section I: Comprehensive Trademark Hit Analysis using chain of thought prompting.
    This approach explicitly walks through the analysis process to ensure consistent results.
    """  
    client = get_azure_client()  
  
    system_prompt = """
    You are a trademark expert attorney specializing in trademark opinion writing. I need you to analyze potential trademark conflicts using chain of thought reasoning.
    
    First, I want you to think step by step:
    
    1. STEP 1 - COORDINATED CLASS ANALYSIS:
    a) Carefully analyze the proposed goods/services: "{goods_services}"
    b) Determine which additional trademark classes are considered related or coordinated with the primary class {class_number}
    c) Provide detailed justification for each coordinated class selected
    d) Produce a finalized list of all trademark classes that should be included in the conflict assessment

    2. STEP 2 - IDENTICAL MARK ANALYSIS:
    a) Identify all trademarks that are an EXACT match to the proposed mark "{mark}" (case-insensitive)
    b) For each identical mark, assess:
        - Is the mark registered in the SAME class as the proposed mark?
        - Is it registered in any of the COORDINATED classes from Step 1?
        - Are the goods/services similar, related, or overlapping with the proposed goods/services?
    c) Clearly specify `class_match` and `goods_services_match` values for each mark
    
    3. STEP 3 - ONE LETTER DIFFERENCE ANALYSIS:
    a) Identify trademarks that differ from the proposed mark by only ONE letter
    b) Acceptable variations include one-letter substitution, addition, or deletion
    c) For each mark, specify the `class_match` and `goods_services_match` values, and document the type of variation

    4. STEP 4 - TWO LETTER DIFFERENCE ANALYSIS:
    a) Identify trademarks that differ from the proposed mark by exactly TWO letters
    b) These may be substitutions, additions, deletions, or a combination thereof
    c) For each mark, specify the `class_match` and `goods_services_match` values, and document the type of variation
 
    STEP 5 – SIMILAR MARK ANALYSIS (Phonetic, Semantic, Functional)
    Perform an analysis to identify trademarks that are similar to the proposed mark. Similarity can be based on:

    1. Phonetic Similarity – Marks that sound similar, even if spelled differently.
    Example: “BRIGHT VISION” is phonetically similar to “HUGE WEIGHT BRIGHT VISIONARY”.
    

    2. Semantic Similarity – Marks with similar meanings.
    Example: “FRESH START” and “NEW BEGINNING”.

    3. Functional Similarity – Marks that serve a similar purpose or convey the same idea/function.
    Example: Two marks used for sugar-free candies with different names.

    Important:
    - Include trademarks not only from the **same Class**, but also from **coordinated Classes** relevant to the proposed mark.
    - The goal is to catch all potential conflicts, even if the classes differ but are related in function or market.

    For each similar mark identified:

    - Clearly explain what makes it similar to the proposed mark.
    - Specify the type of similarity: phonetic, semantic, or functional.
    - Provide a brief reasoning behind the similarity.

    - Return the following for each mark:
    - class_match: true if it is in the same class or a coordinated class, otherwise false
    - goods_services_match: true if the goods/services overlap or are closely related, otherwise false
    
    6. STEP 6 - CROWDED FIELD ANALYSIS:
       a) Calculate the total number of potentially conflicting marks identified
       b) Calculate what percentage of these marks have different owners
       c) Determine if the field is "crowded" (>50% different owners)
       d) Explain the implications for trademark protection
    
    For each conflict you identify, include comprehensive details:
    - The exact mark name
    - The owner name
    - The full goods/services description (not just class numbers)
    - Registration status
    - Class number
    - Whether there's a class match (true/false)
    - Whether there's a goods/services match (true/false)
    
    YOUR RESPONSE MUST BE IN JSON FORMAT:
    {
      "identified_coordinated_classes": [LIST OF RELATED CLASS NUMBERS],
      "coordinated_classes_explanation": "[EXPLANATION OF WHY THESE CLASSES ARE RELATED TO THE PROPOSED TRADEMARK]",
      "identical_marks": [
        {
          "mark": "[TRADEMARK NAME]",
          "owner": "[OWNER NAME]",
          "goods_services": "[GOODS/SERVICES]",
          "status": "[LIVE/DEAD]",
          "class": "[CLASS]",
          "class_match": true|false,
          "goods_services_match": true|false
        }
      ],
      "one_letter_marks": [
        {
          "mark": "[TRADEMARK NAME]",
          "owner": "[OWNER NAME]",
          "goods_services": "[GOODS/SERVICES]",
          "status": "[LIVE/DEAD]",
          "class": "[CLASS]",
          "difference_type": "One Letter",
          "class_match": true|false,
          "goods_services_match": true|false
        }
      ],
      "two_letter_marks": [
        {
          "mark": "[TRADEMARK NAME]",
          "owner": "[OWNER NAME]",
          "goods_services": "[GOODS/SERVICES]",
          "status": "[LIVE/DEAD]",
          "class": "[CLASS]",
          "difference_type": "Two Letter",
          "class_match": true|false,
          "goods_services_match": true|false
        }
      ],
      "similar_marks": [
        {
          "mark": "[TRADEMARK NAME]",
          "owner": "[OWNER NAME]",
          "goods_services": "[GOODS/SERVICES]",
          "status": "[LIVE/DEAD]",
          "class": "[CLASS]",
          "similarity_type": "[Phonetic|Semantic|Functional]",
          "class_match": true|false,
          "goods_services_match": true|false
        }
      ],
      "crowded_field": {
        "is_crowded": true|false,
        "percentage": [PERCENTAGE],
        "explanation": "[EXPLANATION]"
      }
    }
""" 
  
    user_message = f""" 
    Proposed Trademark: {mark}
    Class: {class_number}
    Goods/Services: {goods_services}
    
    Trademark Conflicts:
    {json.dumps(relevant_conflicts, indent=2)}
    
    Analyze ONLY Section I: Comprehensive Trademark Hit Analysis. Walk through each step methodically:
    
    STEP 1: First, carefully analyze the proposed goods/services and identify ALL coordinated classes.
    STEP 2: Then identify EXACT matches to the trademark "{mark}"
    STEP 3: Next, identify marks with ONE letter difference (substitution, addition, or deletion)
    STEP 4: Then identify marks with TWO letter differences
    STEP 5: Finally, identify phonetically, semantically, or functionally similar marks
    STEP 6: Perform crowded field analysis with precise calculations
    
    IMPORTANT REMINDERS:
    - Focus on matches to the ENTIRE trademark name, not just components
    - Include owner names and goods/services details for each mark
    - For Class Match (True/False):
      * First, explicitly identify all coordinated classes related to the proposed goods/services
      * Mark True if the mark's class exactly matches the proposed class "{class_number}"
      * ALSO mark True if the mark's class is in a coordinated or related class grouping you identified
    - For Goods & Services Match (True/False), compare the mark's goods/services to the proposed goods/services
    - Always include the FULL goods/services description in your output, not just the class number
    - For One/Two Letter differences, carefully verify the exact letter count difference
    - For Similar marks, explicitly state whether similarity is Phonetic, Semantic, or Functional
"""  
  
    try:  
        response = client.chat.completions.create(  
            model="gpt-4o",  
            messages=[  
                {"role": "system", "content": system_prompt},  
                {"role": "user", "content": user_message}  
            ],  
            temperature=0.0,  
        )  
  
        if response.choices and len(response.choices) > 0:  
            content = response.choices[0].message.content  
  
            # Extract JSON data  
            json_match = re.search(r'```json\s*(.*?)\s*```|({[\s\S]*})', content, re.DOTALL)  
            if json_match:  
                json_str = json_match.group(1) or json_match.group(2)  
                try:  
                    raw_results = json.loads(json_str)  
                    # Apply consistency checking  
                    corrected_results = consistency_check(mark, raw_results)  
                    return corrected_results  
                except json.JSONDecodeError:  
                    return {  
                        "identified_coordinated_classes": [],
                        "coordinated_classes_explanation": "Unable to identify coordinated classes",
                        "identical_marks": [],  
                        "one_letter_marks": [],  
                        "two_letter_marks": [],  
                        "similar_marks": [],
                        "crowded_field": {
                            "is_crowded": False,
                            "percentage": 0,
                            "explanation": "Unable to determine crowded field status"
                        }
                    }  
            else:  
                return {  
                    "identified_coordinated_classes": [],
                    "coordinated_classes_explanation": "Unable to identify coordinated classes",
                    "identical_marks": [],  
                    "one_letter_marks": [],  
                    "two_letter_marks": [],  
                    "similar_marks": [],
                    "crowded_field": {
                        "is_crowded": False,
                        "percentage": 0,
                        "explanation": "Unable to determine crowded field status"
                    }
                }  
        else:  
            return {  
                "identified_coordinated_classes": [],
                "coordinated_classes_explanation": "Unable to identify coordinated classes",
                "identical_marks": [],  
                "one_letter_marks": [],  
                "two_letter_marks": [],  
                "similar_marks": [],
                "crowded_field": {
                    "is_crowded": False,
                    "percentage": 0,
                    "explanation": "Unable to determine crowded field status"
                }
            }  
    except Exception as e:  
        print(f"Error in section_one_analysis: {str(e)}")  
        return {  
            "identified_coordinated_classes": [],
            "coordinated_classes_explanation": "Error occurred during analysis",
            "identical_marks": [],  
            "one_letter_marks": [],  
            "two_letter_marks": [],  
            "similar_marks": [],
            "crowded_field": {
                "is_crowded": False,
                "percentage": 0,
                "explanation": "Error occurred during analysis"
            }
        }


def component_consistency_check(mark, results):
    """
    Verify component analysis results for consistency and correctness.
    
    Args:
        mark: The proposed trademark
        results: Raw component analysis results
        
    Returns:
        Validated and corrected component analysis results
    """
    corrected_results = results.copy()
    
    # Ensure coordinated classes exist
    if "identified_coordinated_classes" not in corrected_results:
        corrected_results["identified_coordinated_classes"] = []
    
    if "coordinated_classes_explanation" not in corrected_results:
        corrected_results["coordinated_classes_explanation"] = "No coordinated classes identified"
    
    # Check components field
    if "components" not in corrected_results:
        corrected_results["components"] = []
    
    # Validate each component and its marks
    for i, component in enumerate(corrected_results.get("components", [])):
        # Ensure component has name and marks fields
        if "component" not in component:
            component["component"] = f"Component {i+1}"
        
        if "marks" not in component:
            component["marks"] = []
        
        # Ensure component distinctiveness
        if "distinctiveness" not in component:
            # Default to descriptive if not specified
            component["distinctiveness"] = "DESCRIPTIVE"
        
        # Check each mark in the component
        for j, mark_entry in enumerate(component.get("marks", [])):
            # Ensure all required fields exist
            required_fields = ['mark', 'owner', 'goods_services', 'status', 'class', 'class_match', 'goods_services_match']
            for field in required_fields:
                if field not in mark_entry:
                    if field == 'class_match' or field == 'goods_services_match':
                        corrected_results["components"][i]["marks"][j][field] = False
                    else:
                        corrected_results["components"][i]["marks"][j][field] = "Unknown"
    
    # Validate crowded field analysis
    if "crowded_field" not in corrected_results:
        corrected_results["crowded_field"] = {
            "total_hits": 0,
            "distinct_owner_percentage": 0,
            "is_crowded": False,
            "explanation": "Unable to determine crowded field status"
        }
    else:
        # Ensure all required crowded field fields exist
        if "total_hits" not in corrected_results["crowded_field"]:
            corrected_results["crowded_field"]["total_hits"] = 0
            
        if "distinct_owner_percentage" not in corrected_results["crowded_field"]:
            corrected_results["crowded_field"]["distinct_owner_percentage"] = 0
            
        if "is_crowded" not in corrected_results["crowded_field"]:
            corrected_results["crowded_field"]["is_crowded"] = False
            
        if "explanation" not in corrected_results["crowded_field"]:
            corrected_results["crowded_field"]["explanation"] = "Unable to determine crowded field status"
    
    return corrected_results


def section_two_analysis(mark, class_number, goods_services, relevant_conflicts):  
    """Perform Section II: Component Analysis."""  
    client = get_azure_client()  
  
    system_prompt = """
    You are a trademark expert attorney specializing in trademark opinion writing.
    
    Please perform an analysis focusing on Section II: Component Analysis. In this section, you should:

    (a) Identify and break the proposed trademark into its components (if it is compound).
    (b) For each component, analyze marks that incorporate that component.
    (c) For each conflict record, include details such as the owner, goods/services, registration status, and class information.
    (d) Determine flags for both "goods_services_match" and "class_match."

    IMPORTANT INSTRUCTIONS FOR COORDINATED CLASS ANALYSIS:
    • First, analyze the provided goods/services description to identify which trademark classes are closely related or coordinated with the proposed trademark's class.
    • You MUST thoroughly analyze and include conflicts across RELATED and COORDINATED classes, not just exact class matches.
    • Common coordinated class groupings include:
      - Food and beverage products: Consider classes 29, 30, 31, 32, 35, 43
      - Furniture and home goods: Consider classes 20, 35, 42
      - Clothing and fashion: Consider classes 18, 25, 35
      - Technology and software: Consider classes 9, 38, 42
      - Health and beauty: Consider classes 3, 5, 44
      - Entertainment: Consider classes 9, 41, 42
    • However, do not limit yourself to these examples - use your expertise to identify all relevant coordinated classes for the specific goods/services.
    • If ANY component of the proposed trademark appears in ANY other class, this must be flagged.
    • DO NOT MISS conflicts across coordinated classes - this is CRITICAL.

    • When comparing classes, do not only check for an exact match. Explicitly check whether a conflict is registered in a coordinated or related class; for instance, if a conflict is in a related class grouping, you must mark it with class_match = True.
    • Additionally, perform a crowded field analysis by including total compound counts, the percentage of marks from different owners, and a determination of whether the field is crowded.
    • Return your answer in JSON format with keys "components" and "crowded_field."
    
    YOUR RESPONSE MUST BE IN JSON FORMAT:
    {
      "identified_coordinated_classes": [LIST OF RELATED CLASS NUMBERS],
      "coordinated_classes_explanation": "[EXPLANATION OF WHY THESE CLASSES ARE RELATED TO THE PROPOSED TRADEMARK]",
      "components": [
        {
          "component": "[COMPONENT NAME]",
          "marks": [
            {
              "mark": "[TRADEMARK NAME]",
              "owner": "[OWNER NAME]",
              "goods_services": "[GOODS/SERVICES]",
              "status": "[LIVE/DEAD]",
              "class": "[CLASS]",
              "class_match": true|false,
              "goods_services_match": true|false
            }
          ],
          "distinctiveness": "[GENERIC|DESCRIPTIVE|SUGGESTIVE|ARBITRARY|FANCIFUL]"
        }
      ],
      "crowded_field": {
        "total_hits": [NUMBER],
        "distinct_owner_percentage": [PERCENTAGE],
        "is_crowded": true|false,
        "explanation": "[DETAILED EXPLANATION OF FINDINGS, INCLUDING REDUCED RISK IF is_crowded=true]"
      }
    }
"""  
  
    user_message = f"""
    Proposed Trademark: {mark}
    Class: {class_number}
    Goods/Services: {goods_services}
    
    Trademark Conflicts:
    {json.dumps(relevant_conflicts, indent=2)}
    
    Analyze ONLY Section II: Component Analysis.
    
    IMPORTANT REMINDERS:
    - Include exact counts and percentages for all statistics
    - For Crowded Field Analysis:
      1. Show the total number of compound mark hits
      2. Calculate percentage of marks with different owners
      3. If >50% have different owners, set is_crowded=true and mention decreased risk
    - For Class Match (True/False):
      * First, identify all coordinated classes related to the proposed goods/services "{goods_services}"
      * Mark True if the mark's class exactly matches the proposed class "{class_number}"
      * ALSO mark True if the mark's class is in a coordinated or related class grouping you identified
    - For Goods & Services Match (True/False), compare the mark's goods/services to the proposed goods/services "{goods_services}"
    - Always include the FULL goods/services description in your output, not just the class number
"""  
  
    try:  
        response = client.chat.completions.create(  
            model="gpt-4o",  
            messages=[  
                {"role": "system", "content": system_prompt},  
                {"role": "user", "content": user_message}  
            ],  
            temperature=0.0,  
        )  
  
        if response.choices and len(response.choices) > 0:  
            content = response.choices[0].message.content  
  
            # Extract JSON data  
            json_match = re.search(r'```json\s*(.*?)\s*```|({[\s\S]*})', content, re.DOTALL)  
            if json_match:  
                json_str = json_match.group(1) or json_match.group(2)  
                try:  
                    raw_results = json.loads(json_str)
                    # Apply consistency checking
                    corrected_results = component_consistency_check(mark, raw_results)
                    return corrected_results
                except json.JSONDecodeError:  
                    return {
                        "identified_coordinated_classes": [],
                        "coordinated_classes_explanation": "Unable to identify coordinated classes",
                        "components": [],  
                        "crowded_field": {  
                            "total_hits": 0,
                            "distinct_owner_percentage": 0,
                            "is_crowded": False,
                            "explanation": "Unable to determine crowded field status."  
                        }  
                    }  
            else:  
                return {
                    "identified_coordinated_classes": [],
                    "coordinated_classes_explanation": "Unable to identify coordinated classes",
                    "components": [],  
                    "crowded_field": {  
                        "total_hits": 0,
                        "distinct_owner_percentage": 0,
                        "is_crowded": False,
                        "explanation": "Unable to determine crowded field status."  
                    }  
                }  
        else:  
            return {
                "identified_coordinated_classes": [],
                "coordinated_classes_explanation": "Unable to identify coordinated classes",
                "components": [],  
                "crowded_field": {  
                    "total_hits": 0,
                    "distinct_owner_percentage": 0,
                    "is_crowded": False,
                    "explanation": "Unable to determine crowded field status."  
                }  
            }  
    except Exception as e:  
        print(f"Error in section_two_analysis: {str(e)}")  
        return {
            "identified_coordinated_classes": [],
            "coordinated_classes_explanation": "Error occurred during analysis",
            "components": [],  
            "crowded_field": {  
                "total_hits": 0,
                "distinct_owner_percentage": 0,
                "is_crowded": False,
                "explanation": "Error occurred during analysis"  
            }  
        }


def section_three_analysis(mark, class_number, goods_services, section_one_results, section_two_results):
    """
    Perform Section III: Risk Assessment and Summary
    
    Args:
        mark: The proposed trademark
        class_number: The class of the proposed trademark
        goods_services: The goods and services of the proposed trademark
        section_one_results: Results from Section I
        section_two_results: Results from Section II
        
    Returns:
        A structured risk assessment and summary
    """
    client = get_azure_client()
    
    system_prompt = """
    You are a trademark expert attorney specializing in trademark opinion writing.
    
    Please analyze the results from Sections I and II to create Section III: Risk Assessment and Summary. Your analysis should address:

    Likelihood of Confusion – Evaluate the potential for confusion between the proposed trademark and conflicting marks, including the impact of coordinated class conflicts.
    Descriptiveness – Assess whether the proposed trademark's goods/services are descriptive compared to the conflicts.
    Aggressive Enforcement and Litigious Behavior – Identify any patterns of aggressive enforcement among the owners of the conflicting marks.
    Overall Risk – Provide a risk rating (HIGH, MEDIUM-HIGH, MEDIUM, MEDIUM-LOW, LOW) along with an explanation.

    IMPORTANT:
    • In your risk analysis, consider the fact that a conflict may come from a coordinated or related class. For example, if a conflict is registered under a class different from 20 but falls within a related grouping (such as a class that frequently aligns with furniture or home furnishings), mention this and incorporate its impact on risk.
    • Also, include metrics from crowded field analysis to determine if overlapping conflicting marks reduce the overall risk.
    • Return your response in JSON format with keys: likelihood_of_confusion, descriptiveness, aggressive_enforcement, and overall_risk.
    
    YOUR RESPONSE MUST BE IN JSON FORMAT:
    {
      "likelihood_of_confusion": [
        "[KEY POINT ABOUT LIKELIHOOD OF CONFUSION]",
        "[ADDITIONAL POINT ABOUT LIKELIHOOD OF CONFUSION]"
      ],
      "descriptiveness": [
        "[KEY POINT ABOUT DESCRIPTIVENESS]"
      ],
      "aggressive_enforcement": {
        "owners": [
          {
            "name": "[OWNER NAME]",
            "enforcement_patterns": [
              "[PATTERN 1]",
              "[PATTERN 2]"
            ]
          }
        ],
        "enforcement_landscape": [
          "[KEY POINT ABOUT ENFORCEMENT LANDSCAPE]",
          "[ADDITIONAL POINT ABOUT ENFORCEMENT LANDSCAPE]"
        ]
      },
      "overall_risk": {
        "level": "[HIGH|MEDIUM-HIGH|MEDIUM|MEDIUM-LOW|LOW]",
        "explanation": "[EXPLANATION OF RISK LEVEL WITH FOCUS ON CROWDED FIELD]",
        "crowded_field_percentage": [PERCENTAGE]
      }
    }
    """
    
    user_message = f"""
    Proposed Trademark: {mark}
    Class: {class_number}
    Goods and Services: {goods_services}
    
    Section I Results:
    {json.dumps(section_one_results, indent=2)}
    
    Section II Results:
    {json.dumps(section_two_results, indent=2)}
    
    Create Section III: Risk Assessment and Summary.
    
    IMPORTANT REMINDERS:
    - Focus the risk discussion on crowded field analysis
    - Include the percentage of overlapping marks from crowded field analysis
    - Do NOT include recommendations
    - If the risk is Medium-High and a crowded field is identified, reduce it to Medium-Low
    - For aggressive enforcement analysis, examine the owners of similar marks and identify any known for litigious behavior
    - Specifically analyze coordinated class conflicts - marks in related class groupings may present significant risk even if they're not in the exact same class
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            temperature=0.0,
        )
        
        if response.choices and len(response.choices) > 0:
            content = response.choices[0].message.content
            
            # Extract JSON data
            json_match = re.search(r'```json\s*(.*?)\s*```|({[\s\S]*})', content, re.DOTALL)
            if json_match:
                json_str = json_match.group(1) or json_match.group(2)
                try:
                    return json.loads(json_str)
                except json.JSONDecodeError:
                    return {
                        "likelihood_of_confusion": ["Unable to determine likelihood of confusion."],
                        "descriptiveness": ["Unable to determine descriptiveness."],
                        "aggressive_enforcement": {
                            "owners": [],
                            "enforcement_landscape": ["Unable to determine enforcement patterns."]
                        },
                        "overall_risk": {
                            "level": "MEDIUM",
                            "explanation": "Unable to determine precise risk level.",
                            "crowded_field_percentage": 0
                        }
                    }
            else:
                return {
                    "likelihood_of_confusion": ["Unable to determine likelihood of confusion."],
                    "descriptiveness": ["Unable to determine descriptiveness."],
                    "aggressive_enforcement": {
                        "owners": [],
                        "enforcement_landscape": ["Unable to determine enforcement patterns."]
                    },
                    "overall_risk": {
                        "level": "MEDIUM",
                        "explanation": "Unable to determine precise risk level.",
                        "crowded_field_percentage": 0
                    }
                }
        else:
            return {
                "likelihood_of_confusion": ["Unable to determine likelihood of confusion."],
                "descriptiveness": ["Unable to determine descriptiveness."],
                "aggressive_enforcement": {
                    "owners": [],
                    "enforcement_landscape": ["Unable to determine enforcement patterns."]
                },
                "overall_risk": {
                    "level": "MEDIUM",
                    "explanation": "Unable to determine precise risk level.",
                    "crowded_field_percentage": 0
                }
            }
    except Exception as e:
        print(f"Error in section_three_analysis: {str(e)}")
        return {
            "likelihood_of_confusion": ["Unable to determine likelihood of confusion."],
            "descriptiveness": ["Unable to determine descriptiveness."],
            "aggressive_enforcement": {
                "owners": [],
                "enforcement_landscape": ["Unable to determine enforcement patterns."]
            },
            "overall_risk": {
                "level": "MEDIUM",
                "explanation": "Unable to determine precise risk level.",
                "crowded_field_percentage": 0
            }
        }

def generate_trademark_opinion(conflicts_array, proposed_name, proposed_class, proposed_goods_services):
    """
    Generate a comprehensive trademark opinion by running the entire analysis process.
    
    Args:
        conflicts_array: List of potential trademark conflicts
        proposed_name: Name of the proposed trademark
        proposed_class: Class of the proposed trademark
        proposed_goods_services: Goods and services description
        
    Returns:
        A comprehensive trademark opinion
    """
    # Pre-filter trademarks to get the excluded count
    relevant_conflicts, excluded_count = validate_trademark_relevance(conflicts_array, proposed_goods_services)
    
    print("Performing Section I: Comprehensive Trademark Hit Analysis...")
    section_one_results = section_one_analysis(proposed_name, proposed_class, proposed_goods_services, relevant_conflicts)
    
    print("Performing Section II: Component Analysis...")
    section_two_results = section_two_analysis(proposed_name, proposed_class, proposed_goods_services, relevant_conflicts)
    
    print("Performing Section III: Risk Assessment and Summary...")
    section_three_results = section_three_analysis(proposed_name, proposed_class, proposed_goods_services, section_one_results, section_two_results)
    
    # Create a comprehensive opinion structure
    opinion_structure = {
        "proposed_name": proposed_name,
        "proposed_class": proposed_class,
        "proposed_goods_services": proposed_goods_services,
        "excluded_count": excluded_count,
        "section_one": section_one_results,
        "section_two": section_two_results,
        "section_three": section_three_results
    }
    
    # Format the opinion in a structured way
    comprehensive_opinion = f"""
    REFINED TRADEMARK OPINION: {proposed_name}
    Class: {proposed_class}
    Goods and Services: {proposed_goods_services}

    Section I: Comprehensive Trademark Hit Analysis
    
    (a) Identical Marks:
    {json.dumps(section_one_results.get('identical_marks', []), indent=2)}
    
    (b) One Letter and Two Letter Analysis:
    {json.dumps({
        'one_letter_marks': section_one_results.get('one_letter_marks', []),
        'two_letter_marks': section_one_results.get('two_letter_marks', [])
    }, indent=2)}
    
    (c) Phonetically, Semantically & Functionally Similar Analysis:
    {json.dumps(section_one_results.get('similar_marks', []), indent=2)}
    
    (d) Crowded Field Analysis:
    {json.dumps(section_one_results.get('crowded_field', {}), indent=2)}

    Section II: Component Analysis
    
    (a) Component Analysis:
    {json.dumps(section_two_results.get('components', []), indent=2)}
    
    (b) Crowded Field Analysis:
    {json.dumps(section_two_results.get('crowded_field', {}), indent=2)}

    Section III: Risk Assessment and Summary
    
    Likelihood of Confusion:
    {json.dumps(section_three_results.get('likelihood_of_confusion', []), indent=2)}
    
    Descriptiveness:
    {json.dumps(section_three_results.get('descriptiveness', []), indent=2)}
    
    Overall Risk Level:
    {json.dumps(section_three_results.get('overall_risk', {}), indent=2)}
    
    Note: {excluded_count} trademarks with unrelated goods/services were excluded from this analysis.
    """
    
    # Clean and format the final opinion
    print("Cleaning and formatting the final opinion...")
    formatted_opinion = clean_and_format_opinion(comprehensive_opinion, opinion_structure)
    
    return formatted_opinion


# Example usage function
def run_trademark_analysis(proposed_name, proposed_class, proposed_goods_services, conflicts_data):
    """
    Run a complete trademark analysis with proper error handling.
    
    Args:
        proposed_name: Name of the proposed trademark
        proposed_class: Class of the proposed trademark
        proposed_goods_services: Goods and services of the proposed trademark
        conflicts_data: Array of potential conflict trademarks
        
    Returns:
        A comprehensive trademark opinion
    """
    try:
        if not proposed_name or not proposed_class or not proposed_goods_services:
            return "Error: Missing required trademark information."
            
        if not conflicts_data:
            return "Error: No conflict data provided for analysis."
            
        opinion = generate_trademark_opinion(conflicts_data, proposed_name, proposed_class, proposed_goods_services)
        return opinion
        
    except Exception as e:
        return f"Error running trademark analysis: {str(e)}"

# TAMIL CODE END'S HERE ---------------------------------------------------------------------------------------------------------------------------

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_trademark_opinion_to_word(opinion_output):
    """
    Export trademark opinion to Word document with proper formatting and table support
    """
    document = Document()
    
    # Parse and handle different sections
    lines = opinion_output.split('\n')
    for line in lines:
        line = line.strip()
        
        # Handle table rows
        if '|' in line and 'Trademark' not in line and '---' not in line:
            # Split the line into table cells
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            
            # Create table if it doesn't exist
            if not hasattr(document, 'current_table'):
                document.current_table = document.add_table(rows=1, cols=len(cells))
                document.current_table.style = 'Table Grid'
                
                # Add header row
                for i, cell in enumerate(cells):
                    document.current_table.cell(0, i).text = cell
            else:
                # Add data row
                row_cells = document.current_table.add_row().cells
                for i, cell in enumerate(cells):
                    if i < len(row_cells):  # Ensure index is within bounds
                        row_cells[i].text = cell

        # Handle regular paragraphs
        elif line and '|' not in line:
            document.add_paragraph(line)
    
    # Save the document
    filename = "Trademark_Opinion.docx"
    document.save(filename)
    return filename

# ------- 

from typing import List  
import fitz  # PyMuPDF  
from PIL import Image  
import io  
  
  
def Web_CommonLaw_Overview_List(document: str, start_page: int, pdf_document: fitz.Document) -> List[int]:  
    """  
    Extract the page numbers for the 'Web Common Law Overview List' section.  
    """  
    pages_with_overview = []  
    for i in range(start_page, min(start_page + 2, pdf_document.page_count)):  
        page = pdf_document.load_page(i)  
        page_text = page.get_text()  
        if "Record Nr." in page_text:  # Check for "Record Nr." in the text  
            pages_with_overview.append(i + 1)  # Use 1-based indexing for page numbers  
    return pages_with_overview  
  
  
def convert_pages_to_pil_images(pdf_document: fitz.Document, page_numbers: List[int]) -> List[Image.Image]:  
    """  
    Convert the specified pages of the PDF to PIL images and return them as a list of PIL Image objects.  
    """  
    images = []  
    for page_num in page_numbers:  
        page = pdf_document.load_page(page_num - 1)  # Convert 1-based index to 0-based  
        pix = page.get_pixmap()  # Render the page to a pixmap  
        img = Image.open(io.BytesIO(pix.tobytes("png")))  # Convert pixmap to PIL Image  
        images.append(img)  # Add the PIL Image object to the list  
    return images  
  
  
def web_law_page(document_path: str) -> List[Image.Image]:  
    """  
    Return PIL Image objects of the pages where either:  
    1. "Web Common Law Summary Page:" appears, or  
    2. Both "Web Common Law Overview List" and "Record Nr." appear.  
    """  
    matching_pages = []  # List to store matching page numbers  
  
    with fitz.open(document_path) as pdf_document:  
        for page_num in range(pdf_document.page_count):  
            page = pdf_document.load_page(page_num)  
            page_text = page.get_text()  
            print(page_text)  
              
            # Check for "Web Common Law Summary Page:"  
            if "Web Common Law Page:" in page_text:  
                matching_pages.append(page_num + 1)  
  
  
            # Check for "Web Common Law Overview List" and "Record Nr."  
            if "WCL-" in page_text:  
                matching_pages.append(page_num + 1)  
            # if "Web Common Law Overview List" in page_text and "Record Nr." in page_text:  
            #     overview_pages = Web_CommonLaw_Overview_List(  
            #         page_text, page_num, pdf_document  
            #     )  
            #     matching_pages.extend(overview_pages)  
  
  
        # Remove duplicates and sort the page numbers  
        matching_pages = sorted(set(matching_pages))  
  
        # Convert matching pages to PIL images  
        images = convert_pages_to_pil_images(pdf_document, matching_pages)  
  
    return images  
                
# ---- extraction logic

import io  
import base64  
import cv2  
import json  
import requests  
import os
from PIL import Image  
from typing import List  
import numpy as np
  
# Function to encode images using OpenCV  
def encode_image(image: Image.Image) -> str:  
    """  
    Encode a PIL Image as Base64 string using OpenCV.  
    """  
    # Convert PIL Image to numpy array for OpenCV  
    image_np = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)  
    buffered = cv2.imencode(".jpg", image_np)[1]  
    return base64.b64encode(buffered).decode("utf-8")  
  
  
# Function to process a single image and get the response from LLM  
def process_single_image(image: Image.Image, proposed_name: str) -> dict:  
    """  
    Process a single image by sending it to Azure OpenAI API.  
    Cited term: Check for {proposed_name} in the image.
    """        
    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    api_key = os.getenv("AZURE_API_KEY")
    model="gpt-4o"  

    # Encode the image into Base64 using OpenCV  
    base64_image = encode_image(image)  
  
    # Prepare the prompt for the LLM  
    prompt = f"""Extract the following details from the given image: Cited term, Owner name, Goods & services.\n\n
    
                Cited Term:\n
                - This is the snippet in the product/site text that *fully or partially matches* the physically highlighted or searched trademark name: {proposed_name}.
                - You must prioritize any match that closely resembles '{proposed_name}' — e.g., 'ColorGrip', 'COLORGRIP', 'Color self Grip' , 'Grip Colour', 'color-grip', 'Grip' , or minor variations in spacing/punctuation.

                Owner Name (Brand):\n
                - Identify the name of the individual or entity that owns or manufactures the product.
                - Look for indicators like "Owner:," "Brand:," "by:," or "Manufacturer:."
                - If none are found, return "Not specified."
                
                Goods & Services:\n
                - Extract the core goods and services associated with the trademark or product.  
                - Provide relevant detail (e.g., "permanent hair color," "nail care polish," "hair accessories," or "hair styling tools").
    
                Return output only in the exact below-mentioned format:  
                Example output format:  
                    Cited_term: ColourGrip,\n  
                    Owner_name: Matrix, \n 
                    Goods_&_services: Hair color products,\n    
"""
  
    # Prepare the API payload  
    data = {  
        "model": model,  
        "messages": [  
            {  
                "role": "system",  
                "content": "You are a helpful assistant for extracting Meta Data based on the given Images [Note: Only return the required extracted data in the exact format mentioned].",  
            },  
            {  
                "role": "user",  
                "content": [  
                    {"type": "text", "text": prompt},  
                    {  
                        "type": "image_url",  
                        "image_url": {  
                            "url": f"data:image/png;base64,{base64_image}"  
                        },  
                    },  
                ],  
            },  
        ],  
        "max_tokens": 200,  
        "temperature": 0,  
    }  
  
    # Send the API request  
    headers = {"Content-Type": "application/json", "api-key": api_key}  
    response = requests.post(  
        f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version=2024-10-01-preview",  
        headers=headers,  
        data=json.dumps(data),  
    )  
  
    # Parse the response  
    if response.status_code == 200:  
        extracted_data = response.json()["choices"][0]["message"]["content"]  
    else:  
        extracted_data = "Failed to extract data"    
    # Return the extracted data  
    return {extracted_data.strip()}  
  
  
# Function to process all images one by one  
def extract_web_common_law(page_images: List[Image.Image], proposed_name: str) -> List[dict]:  
    """  
    Send images one by one to Azure OpenAI GPT models,  
    and collect the responses into a single array.  
    """    
    # Process each image and collect the results  
    results = []  
    for idx, image in enumerate(page_images):  
        result = process_single_image(image, proposed_name)  
        results.append(result)  
  
    # Return the collected results as a single array  
    return results  

def analyze_web_common_law(extracted_data: List[str], proposed_name: str) -> str:
    """
    Comprehensive analysis of web common law trademark data through three specialized stages.
    Returns a professional opinion formatted according to legal standards.
    """
    # Stage 1: Cited Term Analysis
    cited_term_analysis = perform_cited_term_analysis(extracted_data, proposed_name)
    
    # Stage 2: Component Analysis
    component_analysis = perform_component_analysis(extracted_data, proposed_name)
    
    # Stage 3: Final Risk Assessment
    risk_assessment = perform_risk_assessment(cited_term_analysis, component_analysis, proposed_name)
    
    # Combine all sections into final report
    final_report = f"""
WEB COMMON LAW OPINION: {proposed_name}

{cited_term_analysis}

{component_analysis}

{risk_assessment}
"""
    return final_report

def perform_cited_term_analysis(extracted_data: List[str], proposed_name: str) -> str:
    """
    Perform Section IV: Comprehensive Cited Term Analysis
    """
    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    api_key = os.getenv("AZURE_API_KEY")
    model = "gpt-4o"

    extracted_text = "\n".join([str(item) for item in extracted_data])
    
    prompt = f"""You are a trademark attorney analyzing web common law trademark data.
Perform Section IV analysis (Comprehensive Cited Term Analysis) with these subsections:

1. Identical Cited Terms
2. One Letter and Two Letter Differences
3. Phonetically/Semantically/Functionally Similar Terms

Analyze this web common law data against proposed trademark: {proposed_name}

Extracted Data:
{extracted_text}

Perform comprehensive analysis:
1. Check for identical cited terms
2. Analyze one/two letter differences
3. Identify similar terms (phonetic/semantic/functional)
4. For each, determine if goods/services are similar

Return results in EXACTLY this format:

Section IV: Comprehensive Cited Term Analysis
(a) Identical Cited Terms:
| Cited Term | Owner | Goods & Services | Goods & Services Match |
|------------|--------|------------------|------------------------|
| [Term 1] | [Owner] | [Goods/Services] | [True/False] |

(b) One Letter and Two Letter Analysis:
| Cited Term | Owner | Goods & Services | Difference Type | Goods & Services Match |
|------------|--------|------------------|----------------|------------------------|
| [Term 1] | [Owner] | [Goods/Services] | [One/Two Letter] | [True/False] |

(c) Phonetically, Semantically & Functionally Similar Analysis:
| Cited Term | Owner | Goods & Services | Similarity Type | Goods & Services Match |
|------------|--------|------------------|-----------------|------------------------|
| [Term 1] | [Owner] | [Goods/Services] | [Phonetic/Semantic/Functional] | [True/False] |

Evaluation Guidelines:
- Goods/services match if they overlap with proposed trademark's intended use
- One letter difference = exactly one character changed/added/removed
- Two letter difference = exactly two characters changed/added/removed
- Phonetic similarity = sounds similar when spoken
- Semantic similarity = similar meaning
- Functional similarity = similar purpose/use
- State "None" when no results are found
- Filter out rows where both match criteria are False
- Always include complete goods/services text
"""

    data = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "You are a trademark attorney specializing in comprehensive trademark analysis. Provide precise, professional analysis in the exact requested format.",
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
        "max_tokens": 2000,
        "temperature": 0.1,
    }

    headers = {"Content-Type": "application/json", "api-key": api_key}
    response = requests.post(
        f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version=2024-10-01-preview",
        headers=headers,
        data=json.dumps(data),
    )

    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    return "Failed to generate cited term analysis"

def perform_component_analysis(extracted_data: List[str], proposed_name: str) -> str:
    """
    Perform Section V: Component Analysis and Crowded Field Assessment
    """
    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    api_key = os.getenv("AZURE_API_KEY")
    model = "gpt-4o"

    extracted_text = "\n".join([str(item) for item in extracted_data])
    
    prompt = f"""You are a trademark attorney analyzing web common law components.
Perform Section V analysis (Component Analysis) with these subsections:

1. Component Breakdown
2. Crowded Field Analysis

Analyze this web common law data against proposed trademark: {proposed_name}

Extracted Data:
{extracted_text}

Perform component analysis:
1. Break proposed term into meaningful components
2. For each component, find other terms using that component
3. Perform crowded field analysis
4. Determine goods/services matches

Return results in EXACTLY this format:

Section V: Component Analysis
(a) Component Analysis:

Component 1: [First Component]
| Cited Term | Owner | Goods & Services | Goods & Services Match |
|-----------|--------|------------------|------------------------|
| [Term 1] | [Owner] | [Goods/Services] | [True/False] |

(b) Crowded Field Analysis:
- **Total component hits found**: [NUMBER]
- **Terms with different owners**: [NUMBER] ([PERCENTAGE]%)
- **Crowded Field Status**: [YES/NO]
- **Analysis**: 
  [DETAILED EXPLANATION OF FINDINGS INCLUDING RISK IMPLICATIONS IF FIELD IS CROWDED]

IMPORTANT:
1. Break cited term into meaningful components
2. For each component, find other terms using that component
3. Include FULL goods/services descriptions
4. For crowded field:
   - Calculate percentage of distinct owners
   - Field is crowded if >50% different owners
   - Include detailed explanation
5. Assess distinctiveness for each component

Additional Instructions:
- Goods/services match if they overlap with proposed trademark's intended use
- For crowded field, calculate:
  * Total component hits
  * Percentage with different owners
- Distinctiveness levels:
  * Generic: Common term for the goods/services
  * Descriptive: Describes characteristic/quality
  * Suggestive: Suggests qualities (requires imagination)
  * Arbitrary: Common word unrelated to goods/services
  * Fanciful: Invented word
- State "None" when no results are found
- Filter out rows where both match criteria are False
- Always include complete goods/services text
"""

    data = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "You are a trademark attorney specializing in component and crowded field analysis. Provide precise, professional analysis in the exact requested format.",
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
        "max_tokens": 2000,
        "temperature": 0.1,
    }

    headers = {"Content-Type": "application/json", "api-key": api_key}
    response = requests.post(
        f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version=2024-10-01-preview",
        headers=headers,
        data=json.dumps(data),
    )

    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    return "Failed to generate component analysis"

def perform_risk_assessment(cited_term_analysis: str, component_analysis: str, proposed_name: str) -> str:
    """
    Perform Section VI: Final Risk Assessment combining all findings
    """
    azure_endpoint = os.getenv("AZURE_ENDPOINT")
    api_key = os.getenv("AZURE_API_KEY")
    model = "gpt-4o"

    prompt = f"""You are a senior trademark attorney preparing a final risk assessment.
Combine these analysis sections into a comprehensive risk assessment for: {proposed_name}

Cited Term Analysis:
{cited_term_analysis}

Component Analysis:
{component_analysis}

Prepare Section VI: Web Common Law Risk Assessment with these subsections:

1. Market Presence
2. Enforcement Patterns
3. Risk Category for Use
4. Combined Risk Assessment

Return results in EXACTLY this format:

Section VI: Web Common Law Risk Assessment

Market Presence:
- [KEY POINT ABOUT MARKET PRESENCE]

Enforcement Patterns:
- **Known Aggressive Owners**:
  * [Owner 1]: [Enforcement patterns]

Risk Category for Use:
- **[USE RISK LEVEL: HIGH/MEDIUM/LOW]**
- [EXPLANATION OF USE RISK LEVEL]

III. COMBINED RISK ASSESSMENT

Overall Risk Category:
- **[OVERALL RISK LEVEL: HIGH/MEDIUM-HIGH/MEDIUM/MEDIUM-LOW/LOW]**
- [EXPLANATION INCORPORATING BOTH TRADEMARK AND WEB COMMON LAW FINDINGS]

Guidelines:
1. Base assessment strictly on the provided analysis
2. Do not introduce new findings not in the analysis
3. Maintain professional, legal tone
4. Be specific about risk factors
5. Highlight any particularly concerning findings
"""

    data = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "You are a senior trademark attorney specializing in risk assessment. Provide precise, professional analysis in the exact requested format.",
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
        "max_tokens": 1500,
        "temperature": 0.1,
    }

    headers = {"Content-Type": "application/json", "api-key": api_key}
    response = requests.post(
        f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version=2024-10-01-preview",
        headers=headers,
        data=json.dumps(data),
    )

    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    return "Failed to generate risk assessment"

# -------------------

# Streamlit App
st.title("Trademark Document Parser Version 6.9")

# File upload
uploaded_files = st.sidebar.file_uploader(
    "Choose PDF files", type="pdf", accept_multiple_files=True
)

if uploaded_files:
    if st.sidebar.button("Check Conflicts", key="check_conflicts"):
        total_files = len(uploaded_files)
        progress_bar = st.progress(0)
        # progress_label.text(f"Progress: 0%")  --- Needed to set

        for i, uploaded_file in enumerate(uploaded_files):
            # Save uploaded file to a temporary file path
            temp_file_path = f"temp_{uploaded_file.name}"
            with open(temp_file_path, "wb") as f:
                f.write(uploaded_file.read())

            start_time = time.time()

            sp = True
            proposed_trademark_details = extract_proposed_trademark_details(
                temp_file_path
            )

            if proposed_trademark_details:
                proposed_name = proposed_trademark_details.get(
                    "proposed_trademark_name", "N"
                )
                proposed_class = proposed_trademark_details.get(
                    "proposed_nice_classes_number"
                )
                proposed_goods_services = proposed_trademark_details.get(
                    "proposed_goods_services", "N"
                )
                if proposed_goods_services != "N":
                    with st.expander(
                        f"Proposed Trademark Details for {uploaded_file.name}"
                    ):
                        st.write(f"Proposed Trademark name: {proposed_name}")
                        st.write(f"Proposed class-number: {proposed_class}")
                        st.write(
                            f"Proposed Goods & Services: {proposed_goods_services}"
                        )
                    class_list = list_conversion(proposed_class)
                else:
                    st.write(
                        "______________________________________________________________________________________________________________________________"
                    )
                    st.write(
                        f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}"
                    )
                    st.write(
                        "______________________________________________________________________________________________________________________________"
                    )
                    sp = False
            else:

                proposed_trademark_details = extract_proposed_trademark_details2(
                    temp_file_path
                )

                if proposed_trademark_details:
                    proposed_name = proposed_trademark_details.get(
                        "proposed_trademark_name", "N"
                    )
                    proposed_class = proposed_trademark_details.get(
                        "proposed_nice_classes_number"
                    )
                    proposed_goods_services = proposed_trademark_details.get(
                        "proposed_goods_services", "N"
                    )
                    if proposed_goods_services != "N":
                        with st.expander(
                            f"Proposed Trademark Details for {uploaded_file.name}"
                        ):
                            st.write(f"Proposed Trademark name: {proposed_name}")
                            st.write(f"Proposed class-number: {proposed_class}")
                            st.write(
                                f"Proposed Goods & Services: {proposed_goods_services}"
                            )
                        class_list = list_conversion(proposed_class)
                    else:
                        st.write(
                            "______________________________________________________________________________________________________________________________"
                        )
                        st.write(
                            f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}"
                        )
                        st.write(
                            "______________________________________________________________________________________________________________________________"
                        )
                        sp = False
                else:
                    st.error(
                        f"Unable to extract Proposed Trademark Details for {uploaded_file.name}"
                    )
                    sp = False
                    continue

            if sp:
                progress_bar.progress(25)
                # Initialize AzureChatOpenAI

                # s_time = time.time()

                existing_trademarks = parse_trademark_details(temp_file_path)
                st.write(len(existing_trademarks))
                # for i in range(25,46):
                #     progress_bar.progress(i)


# PRAVEEN WEB COMMON LAW CODE START'S HERE-------------------------------------------------------------------------------------------------------------------------

                # Updated usage in your Streamlit code would look like:
                # !!! Function used extract the web common law pages into images
                full_web_common_law = web_law_page(temp_file_path)                

                progress_bar.progress(50)
                st.success(
                    f"Existing Trademarks Data Extracted Successfully for {uploaded_file.name}!"
                )

                # !!! Function used extract the web common law details from the images using LLM 
                extracted_web_law = extract_web_common_law(full_web_common_law, proposed_name)  

                # New comprehensive analysis
                analysis_result = analyze_web_common_law(extracted_web_law, proposed_name)

                # Display results
                with st.expander("Extracted Web Common Law Data"):
                    st.write(extracted_web_law)

                with st.expander("Trademark Legal Analysis"):
                    st.markdown(analysis_result)  # Using markdown for better formatting

                # extracted_web_law ----- Web common law stored in this variable 

# PRAVEEN WEB COMMON LAW CODE END'S HERE-------------------------------------------------------------------------------------------------------------------------


                # e_time = time.time()
                # elap_time = e_time - s_time
                # elap_time = elap_time // 60
                # st.write(f"Time taken for extraction: {elap_time} mins")

                # e_time = time.time()
                # elap_time = e_time - s_time
                # st.write(f"Time taken: {elap_time} seconds")

                # Display extracted details

                nfiltered_list = []
                unsame_class_list = []

                # Iterate over each JSON element in trademark_name_list
                for json_element in existing_trademarks:
                    class_numbers = json_element["international_class_number"]
                    # Check if any of the class numbers are in class_list
                    if any(number in class_list for number in class_numbers):
                        nfiltered_list.append(json_element)
                    else:
                        unsame_class_list.append(json_element)

                existing_trademarks = nfiltered_list
                existing_trademarks_unsame = unsame_class_list

                high_conflicts = []
                moderate_conflicts = []
                low_conflicts = []
                Name_Matchs = []
                no_conflicts = []

                lt = len(existing_trademarks)

                for existing_trademark in existing_trademarks:
                    conflict = compare_trademarks(
                        existing_trademark,
                        proposed_name,
                        proposed_class,
                        proposed_goods_services,
                    )
                    if conflict is not None:
                        if conflict["conflict_grade"] == "High":
                            high_conflicts.append(conflict)
                        elif conflict["conflict_grade"] == "Moderate":
                            moderate_conflicts.append(conflict)
                        elif conflict["conflict_grade"] == "Low":
                            low_conflicts.append(conflict)
                        else:
                            no_conflicts.append(conflict)

                for existing_trademarks in existing_trademarks_unsame:
                    if existing_trademarks["international_class_number"] != []:
                        conflict = assess_conflict(
                            existing_trademarks,
                            proposed_name,
                            proposed_class,
                            proposed_goods_services,
                        )

                        if conflict["conflict_grade"] == "Name-Match":
                            # conflict_validation = compare_trademarks2(existing_trademarks, proposed_name, proposed_class, proposed_goods_services)
                            # if conflict_validation == "Name-Match":
                            Name_Matchs.append(conflict)
                        else:
                            print("Low")
                            # low_conflicts.append(conflict)

                st.sidebar.write("_________________________________________________")
                st.sidebar.subheader("\n\nConflict Grades : \n")
                st.sidebar.markdown(f"File: {proposed_name}")
                st.sidebar.markdown(
                    f"Total number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(Name_Matchs) + len(low_conflicts)}"
                )
                st.sidebar.markdown(f"3 conditions satisfied:  {len(high_conflicts)}")
                st.sidebar.markdown(f"2 conditions satisfied:  {len(moderate_conflicts)}")
                st.sidebar.markdown(f"Name Match's Conflicts: {len(Name_Matchs)}")
                st.sidebar.markdown(f"1 condition satisfied: {len(low_conflicts)}")
                st.sidebar.write("_________________________________________________")

                document = Document()

                # Set page size to landscape  
                section = document.sections[0]  
                new_width, new_height = section.page_height, section.page_width  
                section.page_width = new_width  
                section.page_height = new_height  

                document.add_heading(
                    f"Trademark Conflict List for {proposed_name} (VERSION - 6.9) :"
                )

                document.add_heading("Dashboard :", level=2)
                # document.add_paragraph(f"\n\nTotal number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(Name_Matchs) + len(low_conflicts)}\n- High Conflicts: {len(high_conflicts)}\n- Moderate Conflicts: {len(moderate_conflicts)}\n- Name Match's Conflicts: {len(Name_Matchs)}\n- Low Conflicts: {len(low_conflicts)}\n")

                # Updated Calculate the number of conflicts
                total_conflicts = (
                    len(high_conflicts)
                    + len(moderate_conflicts)
                    + len(Name_Matchs)
                    + len(low_conflicts)
                )

                # Create a table with 5 rows (including the header) and 2 columns
                table = document.add_table(rows=5, cols=2)

                # Set the table style and customize the borders
                table.style = "TableGrid"

                tbl = table._tbl
                tblBorders = OxmlElement("w:tblBorders")

                for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                    border_element = OxmlElement(f"w:{border}")
                    border_element.set(qn("w:val"), "single")
                    border_element.set(
                        qn("w:sz"), "4"
                    )  # This sets the border size; you can adjust it as needed
                    border_element.set(qn("w:space"), "0")
                    border_element.set(qn("w:color"), "000000")
                    tblBorders.append(border_element)

                tbl.append(tblBorders)

                # Fill the first column with labels
                labels = [
                    "Total number of conflicts:",
                    "- 3 conditions satisfied:",
                    "- 2 conditions satisfied:",
                    "- Name Match's Conflicts:",
                    "- 1 condition satisfied:",
                ]

                # Fill the second column with the conflict numbers
                values = [
                    total_conflicts,
                    len(high_conflicts),
                    len(moderate_conflicts),
                    len(Name_Matchs), 
                    len(low_conflicts),
                ]

                p = document.add_paragraph(" ")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)

                document.add_heading(
                    "Trademark Definitions: ", level=2
                )
                # p = document.add_paragraph(" ")
                # p.paragraph_format.line_spacing = Pt(18)
                p = document.add_paragraph("CONDITION 1: MARK: NAME-BASED SIMILARITY (comprised of Exact Match, Semantically Equivalent, Phonetically Equivalent, Primary position match)")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph("CONDITION 2: CLASS: CLASS OVERLAP")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph("CONDITION 3: GOODS/SERVICES: OVERLAPPING GOODS/SERVICES & TARGET MARKETS")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph("DIRECT HIT: Direct Name hit, regardless of the class")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph(" ")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)


                # Populate the table with the labels and values
                for i in range(5):
                    table.cell(i, 0).text = labels[i]
                    table.cell(i, 1).text = str(values[i])

                    # Set the font size to 10 for both cells
                    for cell in table.row_cells(i):
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)

                if len(high_conflicts) > 0:
                    document.add_heading("Trademarks with 3 conditions satisfied:", level=2)
                    # Create a pandas DataFrame from the JSON list
                    df_high = pd.DataFrame(high_conflicts)
                    df_high = df_high.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_high = document.add_table(
                        df_high.shape[0] + 1, df_high.shape[1]
                    )
                    # Set a predefined table style (with borders)
                    table_high.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_high.columns):
                        table_high.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_high.iterrows():
                        for j, value in enumerate(row):
                            cell = table_high.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                if len(moderate_conflicts) > 0:
                    document.add_heading("Trademarks with 2 conditions satisfied:", level=2)
                    # Create a pandas DataFrame from the JSON list
                    df_moderate = pd.DataFrame(moderate_conflicts)
                    df_moderate = df_moderate.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_moderate = document.add_table(
                        df_moderate.shape[0] + 1, df_moderate.shape[1]
                    )
                    # Set a predefined table style (with borders)
                    table_moderate.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_moderate.columns):
                        table_moderate.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_moderate.iterrows():
                        for j, value in enumerate(row):
                            cell = table_moderate.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                if len(Name_Matchs) > 0:
                    document.add_heading(
                        "Trademarks with Name Match's Conflicts:", level=2
                    )
                    # Create a pandas DataFrame from the JSON list
                    df_Name_Matchs = pd.DataFrame(Name_Matchs)
                    df_Name_Matchs = df_Name_Matchs.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_Name_Matchs = document.add_table(
                        df_Name_Matchs.shape[0] + 1, df_Name_Matchs.shape[1]
                    )
                    # Set a predefined table style (with borders)
                    table_Name_Matchs.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_Name_Matchs.columns):
                        table_Name_Matchs.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_Name_Matchs.iterrows():
                        for j, value in enumerate(row):
                            cell = table_Name_Matchs.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                if len(low_conflicts) > 0:
                    document.add_heading("Trademarks with 1 condition satisfied:", level=2)
                    # Create a pandas DataFrame from the JSON list
                    df_low = pd.DataFrame(low_conflicts)
                    df_low = df_low.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_low = document.add_table(df_low.shape[0] + 1, df_low.shape[1])
                    # Set a predefined table style (with borders)
                    table_low.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_low.columns):
                        table_low.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_low.iterrows():
                        for j, value in enumerate(row):
                            cell = table_low.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                def add_conflict_paragraph(document, conflict):
                    p = document.add_paragraph(
                        f"Trademark Name : {conflict.get('Trademark name', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark Status : {conflict.get('Trademark Status', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark Owner : {conflict.get('Trademark Owner', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark Class Number : {conflict.get('Trademark class Number', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark serial number : {conflict.get('Trademark serial number', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark registration number : {conflict.get('Trademark registration number', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(
                        f"Trademark Design phrase : {conflict.get('Trademark design phrase', 'N/A')}"
                    )
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"{conflict.get('reasoning','N/A')}\n")
                    p.paragraph_format.line_spacing = Pt(18)
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)

                if len(high_conflicts) > 0:
                    document.add_heading(
                        "Explanation: Trademarks with 3 conditions satisfied:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in high_conflicts:
                        add_conflict_paragraph(document, conflict)

                if len(moderate_conflicts) > 0:
                    document.add_heading(
                        "Explanation: Trademarks with 2 conditions satisfied:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in moderate_conflicts:
                        add_conflict_paragraph(document, conflict)

                if len(Name_Matchs) > 0:
                    document.add_heading(
                        "Trademarks with Name Match's Conflicts Reasoning:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in Name_Matchs:
                        add_conflict_paragraph(document, conflict)

                if len(low_conflicts) > 0:
                    document.add_heading(
                        "Explanation: Trademarks with 1 condition satisfied:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in low_conflicts:
                        add_conflict_paragraph(document, conflict)



                def add_conflict_paragraph_to_array(conflict):  
                    result = []  
                    result.append(f"Trademark Name : {conflict.get('Trademark name', 'N/A')}")  
                    result.append(f"Trademark Status : {conflict.get('Trademark Status', 'N/A')}")  
                    result.append(f"Trademark Owner : {conflict.get('Trademark Owner', 'N/A')}")  
                    result.append(f"Trademark Class Number : {conflict.get('Trademark class Number', 'N/A')}")  
                    result.append(f"Trademark serial number : {conflict.get('Trademark serial number', 'N/A')}")  
                    result.append(f"Trademark registration number : {conflict.get('Trademark registration number', 'N/A')}")  
                    result.append(f"Trademark Design phrase : {conflict.get('Trademark design phrase', 'N/A')}")  
                    result.append(" ")  # Blank line for spacing  
                    result.append(f"{conflict.get('reasoning', 'N/A')}\n")  
                    result.append(" ")  # Blank line for spacing  
                    return result  
                
                conflicts_array = []  
                
                if len(high_conflicts) > 0:  
                    conflicts_array.append("Explanation: Trademarks with 3 conditions satisfied:")  
                    conflicts_array.append(" ")  # Blank line for spacing  
                    for conflict in high_conflicts:  
                        conflicts_array.extend(add_conflict_paragraph_to_array(conflict))  
                
                if len(moderate_conflicts) > 0:  
                    conflicts_array.append("Explanation: Trademarks with 2 conditions satisfied:")  
                    conflicts_array.append(" ")  # Blank line for spacing  
                    for conflict in moderate_conflicts:  
                        conflicts_array.extend(add_conflict_paragraph_to_array(conflict))  
                
                if len(Name_Matchs) > 0:  
                    conflicts_array.append("Trademarks with Name Match's Conflicts Reasoning:")  
                    conflicts_array.append(" ")  # Blank line for spacing  
                    for conflict in Name_Matchs:  
                        conflicts_array.extend(add_conflict_paragraph_to_array(conflict))  
                
                if len(low_conflicts) > 0:  
                    conflicts_array.append("Explanation: Trademarks with 1 condition satisfied:")  
                    conflicts_array.append(" ")  # Blank line for spacing  
                    for conflict in low_conflicts:  
                        conflicts_array.extend(add_conflict_paragraph_to_array(conflict))  
                    

                # for i in range(70,96):
                #     progress_bar.progress(i)


                progress_bar.progress(100)

                filename = proposed_name
                doc_stream = BytesIO()
                document.save(doc_stream)
                doc_stream.seek(0)
                download_table = f'<a href="data:application/octet-stream;base64,{base64.b64encode(doc_stream.read()).decode()}" download="{filename + " Trademark Conflict Report"}.docx">Download: {filename}</a>'
                st.sidebar.markdown(download_table, unsafe_allow_html=True)
                st.success(
                    f"{proposed_name} Document conflict report successfully completed!"
                )
                
                opinion_output = run_trademark_analysis(proposed_name, proposed_class, proposed_goods_services, conflicts_array)
                st.write("------------------------------------------------------------------------------------------------------------------------------")
                st.write(opinion_output)

                # Export to Word
                filename = export_trademark_opinion_to_word(opinion_output)
                
                # Download button
                with open(filename, "rb") as file:
                    st.sidebar.download_button(
                        label="Download Trademark Opinion",
                        data=file,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                end_time = time.time()
                elapsed_time = end_time - start_time
                elapsed_time = elapsed_time // 60
                st.write(f"Time taken: {elapsed_time} mins")

                st.write(
                    "______________________________________________________________________________________________________________________________"
                )

        progress_bar.progress(100)
        st.success("All documents processed successfully!")
